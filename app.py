from __future__ import annotations

import os
import re
import shutil
import sqlite3
import subprocess
import sys
import webbrowser
import zipfile
from datetime import datetime
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
from tkinter.scrolledtext import ScrolledText
import xml.etree.ElementTree as ET

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT

    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

APP_TITLE = "Tumaini Academy Learner Management System"
BASE_DIR = Path(__file__).resolve().parent
APP_DIR = Path(os.getenv("APPDATA", str(Path.home()))) / "TumainiAcademyLMS"
DB_PATH = APP_DIR / "tumaini_academy.db"
DEFAULT_REPORT_DIR = APP_DIR / "Reports"
DEFAULT_LOGO_PATH = BASE_DIR / "assets" / "tumaini_logo.png"
DEFAULT_IMPORT_XLSX_PATH = Path(r"c:\Users\user\Documents\TUM\Tumaini\tumaini lerners.xlsx")
DATE_FORMAT = "%Y-%m-%d"
TIME_FORMAT = "%H:%M"
PHONE_REGEX = re.compile(r"^\+?[0-9]{10,15}$")
VALID_BOARDING = ("Boarder", "Day Scholar")
DAY_SCHOLAR_TRANSPORT = ("School Bus", "Bicycle", "Walking")
CLASS_LEVELS = (
    "PRE PRIMARY ONE",
    "PRE PRIMARY TWO",
    "GRADE 1",
    "GRADE 2",
    "GRADE 3",
    "GRADE 4",
    "GRADE 5",
    "GRADE 6",
    "GRADE 7",
    "GRADE 8",
    "GRADE 9",
    "GRADE 10",
)
ALL_CLASSES_FILTER = "All Classes"
ALL_BOARDING_FILTER = "All Boarding"
ALL_TRANSPORT_FILTER = "All Transport"
MANUAL_PDF = BASE_DIR / "docs" / "User_Manual.pdf"
MANUAL_MD = BASE_DIR / "docs" / "User_Manual.md"
REPORT_ORIENTATIONS = ("Auto", "Portrait", "Landscape")


def ensure_app_directories() -> None:
    APP_DIR.mkdir(parents=True, exist_ok=True)
    DEFAULT_REPORT_DIR.mkdir(parents=True, exist_ok=True)


def parse_date(date_text: str) -> datetime:
    return datetime.strptime(date_text, DATE_FORMAT)


def is_valid_time(time_text: str) -> bool:
    try:
        datetime.strptime(time_text, TIME_FORMAT)
        return True
    except ValueError:
        return False


def normalize_phone(phone: str) -> str:
    return phone.replace(" ", "").replace("-", "")


def open_path(path: Path) -> None:
    if sys.platform.startswith("win"):
        os.startfile(str(path))
        return
    try:
        subprocess.Popen(["xdg-open", str(path)])
    except Exception:
        subprocess.Popen(["open", str(path)])


def now_date() -> str:
    return datetime.now().strftime(DATE_FORMAT)


def now_time() -> str:
    return datetime.now().strftime(TIME_FORMAT)


def report_timestamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def add_number_column(rows: list[list[str]]) -> tuple[list[str], list[list[str]]]:
    numbered_rows = []
    for index, row in enumerate(rows, start=1):
        numbered_rows.append([str(index)] + [str(item) for item in row])
    return ["No."], numbered_rows


def normalize_admission_no(value: str) -> str:
    return re.sub(r"\s+", "", value.strip()).upper()


def sanitize_phone_for_storage(raw_phone: str) -> str:
    cleaned = normalize_phone(raw_phone.strip())
    cleaned = re.sub(r"[^\d+]", "", cleaned)
    if cleaned.startswith("00"):
        cleaned = f"+{cleaned[2:]}"
    if PHONE_REGEX.match(cleaned):
        return cleaned
    digits = re.sub(r"\D", "", raw_phone)
    if len(digits) == 9:
        digits = f"0{digits}"
    if 10 <= len(digits) <= 15:
        return digits
    return "0000000000"


def normalize_class_level(raw_grade: str) -> str:
    normalized = re.sub(r"\s+", "", raw_grade.strip().upper())
    mapping = {
        "PP1": "PRE PRIMARY ONE",
        "PREPRIMARYONE": "PRE PRIMARY ONE",
        "PP2": "PRE PRIMARY TWO",
        "PREPRIMARYTWO": "PRE PRIMARY TWO",
    }
    if normalized in mapping:
        return mapping[normalized]
    if normalized.startswith("GRADE"):
        normalized = normalized.replace("GRADE", "")
    if normalized.startswith("G"):
        normalized = normalized[1:]
    if normalized.isdigit():
        number = int(normalized)
        if 1 <= number <= 10:
            return f"GRADE {number}"
    return ""


def _excel_col_to_index(cell_ref: str) -> int:
    letters = "".join(ch for ch in cell_ref if ch.isalpha()).upper()
    if not letters:
        return 0
    value = 0
    for ch in letters:
        value = value * 26 + (ord(ch) - ord("A") + 1)
    return value - 1


def _read_xlsx_cell_text(cell: ET.Element, shared_strings: list[str], ns: dict[str, str]) -> str:
    cell_type = cell.attrib.get("t")
    value_node = cell.find("s:v", ns)
    if value_node is None:
        inline_text = cell.find("s:is/s:t", ns)
        return (inline_text.text or "").strip() if inline_text is not None else ""
    raw_value = (value_node.text or "").strip()
    if cell_type == "s":
        if raw_value.isdigit():
            index = int(raw_value)
            if 0 <= index < len(shared_strings):
                return shared_strings[index].strip()
    return raw_value.strip()


def read_learners_from_xlsx(file_path: Path) -> list[dict[str, str]]:
    ns = {
        "s": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
    }
    header_aliases = {
        "adm": "admission_no",
        "admissionno": "admission_no",
        "admissionnumber": "admission_no",
        "name": "learner_name",
        "learnername": "learner_name",
        "grade": "class_level",
        "class": "class_level",
        "parent": "parent_name",
        "parentname": "parent_name",
        "telphoneno": "parent_phone",
        "telephoneno": "parent_phone",
        "phone": "parent_phone",
        "parentphone": "parent_phone",
    }

    with zipfile.ZipFile(file_path, "r") as zf:
        if "xl/workbook.xml" not in zf.namelist():
            raise RuntimeError("Invalid .xlsx file: workbook.xml is missing.")

        shared_strings: list[str] = []
        if "xl/sharedStrings.xml" in zf.namelist():
            shared_root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
            for item in shared_root.findall("s:si", ns):
                parts = [t.text or "" for t in item.findall(".//s:t", ns)]
                shared_strings.append("".join(parts))

        workbook_root = ET.fromstring(zf.read("xl/workbook.xml"))
        rels_root = ET.fromstring(zf.read("xl/_rels/workbook.xml.rels"))
        rel_map = {
            rel.attrib["Id"]: rel.attrib["Target"]
            for rel in rels_root.findall("rel:Relationship", ns)
        }

        first_sheet = workbook_root.find("s:sheets/s:sheet", ns)
        if first_sheet is None:
            raise RuntimeError("No worksheet found in the workbook.")
        rid = first_sheet.attrib.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", ""
        )
        target = rel_map.get(rid, "")
        if not target:
            raise RuntimeError("Could not resolve worksheet target in workbook relationships.")
        if not target.startswith("xl/"):
            target = f"xl/{target}"
        if target not in zf.namelist():
            raise RuntimeError(f"Worksheet file not found inside workbook: {target}")

        worksheet_root = ET.fromstring(zf.read(target))
        xml_rows = worksheet_root.findall("s:sheetData/s:row", ns)
        if not xml_rows:
            return []

        parsed_rows: list[dict[int, str]] = []
        max_col_index = 0
        for xml_row in xml_rows:
            row_map: dict[int, str] = {}
            for cell in xml_row.findall("s:c", ns):
                col_index = _excel_col_to_index(cell.attrib.get("r", "A1"))
                row_map[col_index] = _read_xlsx_cell_text(cell, shared_strings, ns)
                max_col_index = max(max_col_index, col_index)
            parsed_rows.append(row_map)

        if not parsed_rows:
            return []
        header_map = parsed_rows[0]
        headers: dict[int, str] = {}
        for col_idx in range(max_col_index + 1):
            raw_header = header_map.get(col_idx, "")
            normalized_header = re.sub(r"[^a-z0-9]", "", raw_header.lower())
            canonical = header_aliases.get(normalized_header, "")
            if canonical:
                headers[col_idx] = canonical

        required = {"admission_no", "learner_name", "class_level", "parent_name", "parent_phone"}
        if not required.issubset(set(headers.values())):
            raise RuntimeError(
                "Required columns not found. Expected ADM/NAME/GRADE/parent/TELPHONE NO (or equivalents)."
            )

        records: list[dict[str, str]] = []
        for row_map in parsed_rows[1:]:
            record = {
                "admission_no": "",
                "learner_name": "",
                "class_level": "",
                "parent_name": "",
                "parent_phone": "",
            }
            for col_idx, key in headers.items():
                record[key] = row_map.get(col_idx, "").strip()
            if any(record.values()):
                records.append(record)
        return records


class DatabaseManager:
    def __init__(self, db_path: Path) -> None:
        self.db_path = db_path

    def _connect(self) -> sqlite3.Connection:
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        conn.execute("PRAGMA foreign_keys = ON")
        return conn

    def initialize(self) -> None:
        with self._connect() as conn:
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS learners (
                    admission_no TEXT PRIMARY KEY,
                    learner_name TEXT NOT NULL,
                    class_level TEXT NOT NULL DEFAULT 'GRADE 1',
                    parent_name TEXT NOT NULL,
                    parent_phone TEXT NOT NULL,
                    boarding_status TEXT NOT NULL
                        CHECK(boarding_status IN ('Boarder', 'Day Scholar')),
                    transport_mode TEXT NOT NULL
                        CHECK(transport_mode IN ('School Bus', 'Bicycle', 'Walking', 'N/A'))
                )
                """
            )
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS reporting_records (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    admission_no TEXT NOT NULL,
                    learner_name TEXT NOT NULL,
                    report_date TEXT NOT NULL,
                    report_time TEXT NOT NULL,
                    arrival_transport_mode TEXT NOT NULL DEFAULT 'N/A',
                    accompanied_source TEXT NOT NULL DEFAULT 'Registered Parent',
                    accompanied_by_name TEXT NOT NULL DEFAULT '',
                    accompanied_by_phone TEXT NOT NULL DEFAULT '',
                    company TEXT NOT NULL,
                    FOREIGN KEY (admission_no) REFERENCES learners(admission_no) ON DELETE CASCADE
                )
                """
            )
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS settings (
                    setting_key TEXT PRIMARY KEY,
                    setting_value TEXT NOT NULL
                )
                """
            )
            conn.execute(
                "INSERT OR IGNORE INTO settings(setting_key, setting_value) VALUES(?, ?)",
                ("school_name", "Tumaini Academy"),
            )
            conn.execute(
                "INSERT OR IGNORE INTO settings(setting_key, setting_value) VALUES(?, ?)",
                ("default_report_dir", str(DEFAULT_REPORT_DIR)),
            )
            conn.execute(
                "INSERT OR IGNORE INTO settings(setting_key, setting_value) VALUES(?, ?)",
                ("app_password", ""),
            )
            conn.execute(
                "INSERT OR IGNORE INTO settings(setting_key, setting_value) VALUES(?, ?)",
                ("logo_path", str(DEFAULT_LOGO_PATH) if DEFAULT_LOGO_PATH.exists() else ""),
            )
            conn.execute(
                "INSERT OR IGNORE INTO settings(setting_key, setting_value) VALUES(?, ?)",
                ("report_orientation", "Auto"),
            )
            conn.execute(
                "INSERT OR IGNORE INTO settings(setting_key, setting_value) VALUES(?, ?)",
                ("word_template_portrait", ""),
            )
            conn.execute(
                "INSERT OR IGNORE INTO settings(setting_key, setting_value) VALUES(?, ?)",
                ("word_template_landscape", ""),
            )
            conn.execute(
                "INSERT OR IGNORE INTO settings(setting_key, setting_value) VALUES(?, ?)",
                ("pdf_template_portrait", ""),
            )
            conn.execute(
                "INSERT OR IGNORE INTO settings(setting_key, setting_value) VALUES(?, ?)",
                ("pdf_template_landscape", ""),
            )

            self._ensure_column(
                conn,
                "learners",
                "class_level",
                "TEXT NOT NULL DEFAULT 'GRADE 1'",
            )
            self._ensure_column(
                conn,
                "reporting_records",
                "arrival_transport_mode",
                "TEXT NOT NULL DEFAULT 'N/A'",
            )
            self._ensure_column(
                conn,
                "reporting_records",
                "accompanied_source",
                "TEXT NOT NULL DEFAULT 'Registered Parent'",
            )
            self._ensure_column(
                conn,
                "reporting_records",
                "accompanied_by_name",
                "TEXT NOT NULL DEFAULT ''",
            )
            self._ensure_column(
                conn,
                "reporting_records",
                "accompanied_by_phone",
                "TEXT NOT NULL DEFAULT ''",
            )
            conn.execute("DROP INDEX IF EXISTS ux_reporting_adm_date")
            conn.commit()

    def _get_columns(self, conn: sqlite3.Connection, table_name: str) -> set[str]:
        return {
            row["name"]
            for row in conn.execute(f"PRAGMA table_info({table_name})").fetchall()
        }

    def _ensure_column(
        self,
        conn: sqlite3.Connection,
        table_name: str,
        column_name: str,
        definition_sql: str,
    ) -> None:
        columns = self._get_columns(conn, table_name)
        if column_name not in columns:
            conn.execute(f"ALTER TABLE {table_name} ADD COLUMN {column_name} {definition_sql}")

    def validate_database_file(self, db_file: Path) -> bool:
        try:
            with sqlite3.connect(db_file) as conn:
                conn.execute("SELECT 1")
            return True
        except sqlite3.Error:
            return False

    def backup_database(self, destination: Path) -> None:
        shutil.copy2(self.db_path, destination)

    def restore_database(self, source: Path) -> None:
        if not source.exists():
            raise FileNotFoundError(f"Backup file not found: {source}")
        if not self.validate_database_file(source):
            raise RuntimeError("Selected file is not a valid SQLite database.")
        shutil.copy2(source, self.db_path)
    def get_setting(self, key: str, default: str = "") -> str:
        with self._connect() as conn:
            row = conn.execute(
                "SELECT setting_value FROM settings WHERE setting_key = ?",
                (key,),
            ).fetchone()
            return row["setting_value"] if row else default

    def set_setting(self, key: str, value: str) -> None:
        with self._connect() as conn:
            conn.execute(
                """
                INSERT INTO settings(setting_key, setting_value)
                VALUES(?, ?)
                ON CONFLICT(setting_key) DO UPDATE SET setting_value = excluded.setting_value
                """,
                (key, value),
            )
            conn.commit()

    def add_learner(self, learner: dict[str, str]) -> tuple[bool, str]:
        try:
            with self._connect() as conn:
                conn.execute(
                    """
                    INSERT INTO learners (
                        admission_no, learner_name, class_level, parent_name, parent_phone, boarding_status, transport_mode
                    ) VALUES (?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        learner["admission_no"],
                        learner["learner_name"],
                        learner["class_level"],
                        learner["parent_name"],
                        learner["parent_phone"],
                        learner["boarding_status"],
                        learner["transport_mode"],
                    ),
                )
                conn.commit()
            return True, "Learner saved successfully."
        except sqlite3.IntegrityError:
            return False, "Admission number already exists."

    def update_learner(self, learner: dict[str, str]) -> tuple[bool, str]:
        with self._connect() as conn:
            cursor = conn.execute(
                """
                UPDATE learners
                SET learner_name = ?, class_level = ?, parent_name = ?, parent_phone = ?,
                    boarding_status = ?, transport_mode = ?
                WHERE admission_no = ?
                """,
                (
                    learner["learner_name"],
                    learner["class_level"],
                    learner["parent_name"],
                    learner["parent_phone"],
                    learner["boarding_status"],
                    learner["transport_mode"],
                    learner["admission_no"],
                ),
            )
            conn.commit()
            if cursor.rowcount == 0:
                return False, "Learner not found."
            return True, "Learner updated successfully."

    def delete_learner(self, admission_no: str) -> tuple[bool, str]:
        with self._connect() as conn:
            cursor = conn.execute(
                "DELETE FROM learners WHERE admission_no = ?",
                (admission_no,),
            )
            conn.commit()
            if cursor.rowcount == 0:
                return False, "Learner not found."
            return True, "Learner deleted successfully."

    def get_learner(self, admission_no: str) -> sqlite3.Row | None:
        with self._connect() as conn:
            return conn.execute(
                "SELECT * FROM learners WHERE admission_no = ?",
                (admission_no,),
            ).fetchone()

    def search_learners(self, term: str = "") -> list[sqlite3.Row]:
        wildcard = f"%{term.strip()}%"
        with self._connect() as conn:
            return conn.execute(
                """
                SELECT admission_no, learner_name, class_level, parent_name, parent_phone, boarding_status, transport_mode
                FROM learners
                WHERE admission_no LIKE ?
                   OR learner_name LIKE ?
                   OR class_level LIKE ?
                   OR parent_name LIKE ?
                   OR parent_phone LIKE ?
                ORDER BY learner_name ASC
                """,
                (wildcard, wildcard, wildcard, wildcard, wildcard),
            ).fetchall()

    def add_reporting_record(
        self,
        admission_no: str,
        learner_name: str,
        report_date: str,
        report_time: str,
        arrival_transport_mode: str,
        accompanied_source: str,
        accompanied_by_name: str,
        accompanied_by_phone: str,
        company: str,
    ) -> tuple[bool, str]:
        try:
            with self._connect() as conn:
                conn.execute(
                    """
                    INSERT INTO reporting_records(
                        admission_no,
                        learner_name,
                        report_date,
                        report_time,
                        arrival_transport_mode,
                        accompanied_source,
                        accompanied_by_name,
                        accompanied_by_phone,
                        company
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        admission_no,
                        learner_name,
                        report_date,
                        report_time,
                        arrival_transport_mode,
                        accompanied_source,
                        accompanied_by_name,
                        accompanied_by_phone,
                        company,
                    ),
                )
                conn.commit()
            return True, "Arrival reported successfully."
        except sqlite3.Error as error:
            return False, f"Could not save reporting record: {error}"

    def delete_reporting_record(self, record_id: int) -> tuple[bool, str]:
        with self._connect() as conn:
            cursor = conn.execute(
                "DELETE FROM reporting_records WHERE id = ?",
                (record_id,),
            )
            conn.commit()
            if cursor.rowcount == 0:
                return False, "Reporting record not found."
            return True, "Reporting record deleted successfully."

    def get_reporting_for_date(self, report_date: str, term: str = "") -> list[sqlite3.Row]:
        wildcard = f"%{term.strip()}%"
        with self._connect() as conn:
            return conn.execute(
                """
                SELECT
                    id,
                    report_time,
                    learner_name,
                    admission_no,
                    arrival_transport_mode,
                    accompanied_source,
                    accompanied_by_name,
                    accompanied_by_phone,
                    company
                FROM reporting_records
                WHERE report_date = ?
                  AND (
                        learner_name LIKE ?
                     OR admission_no LIKE ?
                     OR company LIKE ?
                     OR accompanied_by_name LIKE ?
                     OR accompanied_by_phone LIKE ?
                  )
                ORDER BY report_time ASC
                """,
                (report_date, wildcard, wildcard, wildcard, wildcard, wildcard),
            ).fetchall()

    def report_reported_learners(self, from_date: str, to_date: str) -> list[sqlite3.Row]:
        with self._connect() as conn:
            return conn.execute(
                """
                SELECT
                    rr.report_date,
                    rr.report_time,
                    rr.admission_no,
                    rr.learner_name,
                    l.class_level,
                    l.boarding_status,
                    rr.arrival_transport_mode,
                    rr.accompanied_by_name,
                    rr.accompanied_by_phone
                FROM reporting_records rr
                INNER JOIN learners l ON rr.admission_no = l.admission_no
                WHERE rr.report_date BETWEEN ? AND ?
                ORDER BY rr.report_date ASC, rr.report_time ASC, rr.learner_name ASC
                """,
                (from_date, to_date),
            ).fetchall()

    def report_boarders(self) -> list[sqlite3.Row]:
        with self._connect() as conn:
            return conn.execute(
                """
                SELECT
                    admission_no,
                    learner_name,
                    class_level,
                    parent_name,
                    parent_phone
                FROM learners
                WHERE boarding_status = 'Boarder'
                ORDER BY learner_name ASC
                """
            ).fetchall()

    def report_school_bus_users(self) -> list[sqlite3.Row]:
        with self._connect() as conn:
            return conn.execute(
                """
                SELECT admission_no, learner_name, parent_name, parent_phone
                FROM learners
                WHERE boarding_status = 'Day Scholar'
                  AND transport_mode = 'School Bus'
                ORDER BY learner_name ASC
                """
            ).fetchall()

    def report_all_learners(self) -> list[sqlite3.Row]:
        with self._connect() as conn:
            return conn.execute(
                """
                SELECT admission_no, learner_name, class_level, parent_name, parent_phone, boarding_status, transport_mode
                FROM learners
                ORDER BY learner_name ASC
                """
            ).fetchall()

    def report_parents_with_multiple_children(self) -> list[sqlite3.Row]:
        with self._connect() as conn:
            return conn.execute(
                """
                SELECT
                    parent_name,
                    parent_phone,
                    COUNT(*) AS child_count,
                    GROUP_CONCAT(learner_name || ' (' || admission_no || ')', ', ') AS children
                FROM learners
                GROUP BY parent_phone
                HAVING COUNT(*) > 1
                ORDER BY child_count DESC, parent_name ASC
                """
            ).fetchall()


class ReportExporter:
    def __init__(self, school_name: str) -> None:
        self.school_name = school_name

    def _resolve_orientation(
        self,
        orientation_preference: str,
        columns: list[str],
        rows: list[list[str]],
    ) -> str:
        pref = (orientation_preference or "Auto").strip().title()
        if pref in ("Portrait", "Landscape"):
            return pref

        if len(columns) >= 8:
            return "Landscape"
        for row in rows[:50]:
            for value in row:
                if len(str(value)) > 32:
                    return "Landscape"
        return "Portrait"

    def export_pdf(
        self,
        title: str,
        columns: list[str],
        rows: list[list[str]],
        output_file: Path,
        from_date: str,
        to_date: str,
        orientation_preference: str = "Auto",
        portrait_template: str = "",
        landscape_template: str = "",
    ) -> str:
        if not REPORTLAB_AVAILABLE:
            raise RuntimeError("PDF export requires reportlab. Install it from requirements.txt.")

        orientation = self._resolve_orientation(orientation_preference, columns, rows)
        page_size = landscape(A4) if orientation == "Landscape" else A4
        doc = SimpleDocTemplate(str(output_file), pagesize=page_size)
        styles = getSampleStyleSheet()
        content: list[object] = []
        content.append(Paragraph(self.school_name, styles["Title"]))
        content.append(Paragraph(title, styles["Heading2"]))
        content.append(Paragraph(f"Generated: {datetime.now():%Y-%m-%d %H:%M:%S}", styles["Normal"]))
        if from_date and to_date:
            content.append(Paragraph(f"Date range: {from_date} to {to_date}", styles["Normal"]))
        content.append(Paragraph(f"Total records: {len(rows)}", styles["Normal"]))
        content.append(Spacer(1, 12))

        table_rows = [columns]
        if rows:
            table_rows.extend(rows)
        else:
            table_rows.append(["No data for this selection."] + [""] * (len(columns) - 1))

        table = Table(table_rows, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1C5D63")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
                ]
            )
        )
        content.append(table)

        template_path_text = landscape_template if orientation == "Landscape" else portrait_template
        template_path = Path(template_path_text) if template_path_text else None

        if template_path and template_path.exists() and template_path.suffix.lower() in (".png", ".jpg", ".jpeg"):
            def _draw_template(canvas_obj, doc_obj) -> None:
                page_w, page_h = doc_obj.pagesize
                canvas_obj.saveState()
                canvas_obj.drawImage(
                    str(template_path),
                    0,
                    0,
                    width=page_w,
                    height=page_h,
                    preserveAspectRatio=False,
                    mask="auto",
                )
                canvas_obj.restoreState()

            doc.build(content, onFirstPage=_draw_template, onLaterPages=_draw_template)
        else:
            doc.build(content)
        return orientation

    def export_docx(
        self,
        title: str,
        columns: list[str],
        rows: list[list[str]],
        output_file: Path,
        from_date: str,
        to_date: str,
        orientation_preference: str = "Auto",
        portrait_template: str = "",
        landscape_template: str = "",
    ) -> str:
        if not DOCX_AVAILABLE:
            raise RuntimeError("Word export requires python-docx. Install it from requirements.txt.")

        orientation = self._resolve_orientation(orientation_preference, columns, rows)
        template_path_text = landscape_template if orientation == "Landscape" else portrait_template
        template_path = Path(template_path_text) if template_path_text else None

        if template_path and template_path.exists() and template_path.suffix.lower() == ".docx":
            doc = Document(str(template_path))
        else:
            doc = Document()

        for section in doc.sections:
            if orientation == "Landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                if section.page_width < section.page_height:
                    section.page_width, section.page_height = section.page_height, section.page_width
            else:
                section.orientation = WD_ORIENT.PORTRAIT
                if section.page_width > section.page_height:
                    section.page_width, section.page_height = section.page_height, section.page_width

        doc.add_heading(self.school_name, 0)
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"Generated: {datetime.now():%Y-%m-%d %H:%M:%S}")
        if from_date and to_date:
            doc.add_paragraph(f"Date range: {from_date} to {to_date}")
        doc.add_paragraph(f"Total records: {len(rows)}")

        if rows:
            table = doc.add_table(rows=1, cols=len(columns))
            table.style = "Table Grid"
            for index, col_name in enumerate(columns):
                table.rows[0].cells[index].text = col_name
            for row in rows:
                cells = table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = str(value)
        else:
            doc.add_paragraph("No data for this selection.")
        doc.save(output_file)
        return orientation


class TumainiApp(tk.Tk):
    def __init__(self) -> None:
        super().__init__()
        ensure_app_directories()
        self.db = DatabaseManager(DB_PATH)
        self.db.initialize()

        self.school_name = self.db.get_setting("school_name", "Tumaini Academy")
        self.default_report_dir = Path(
            self.db.get_setting("default_report_dir", str(DEFAULT_REPORT_DIR))
        )
        self.default_report_dir.mkdir(parents=True, exist_ok=True)
        saved_logo = self.db.get_setting("logo_path", "")
        self.logo_path = Path(saved_logo) if saved_logo else DEFAULT_LOGO_PATH
        self.logo_image: tk.PhotoImage | None = None
        self.last_generated_report: Path | None = None
        self.selected_reporting_adm: str | None = None
        self.selected_reporting_parent_name = ""
        self.selected_reporting_parent_phone = ""
        self.selected_reporting_status = ""
        self.selected_reporting_transport = ""

        self.title(APP_TITLE)
        self.geometry("1220x780")
        self.minsize(1100, 720)

        self._build_variables()
        self._build_menu()
        self._build_layout()
        self._load_initial_data()
        self.protocol("WM_DELETE_WINDOW", self._exit_app)
        self._active_mousewheel_canvas: tk.Canvas | None = None
        self.bind_all("<MouseWheel>", self._on_global_mousewheel, add="+")
        self.bind_all("<Button-4>", self._on_global_mousewheel, add="+")
        self.bind_all("<Button-5>", self._on_global_mousewheel, add="+")

        if not self._check_optional_login():
            self.after(20, self.destroy)

    def _build_variables(self) -> None:
        self.status_text = tk.StringVar(value="Ready")

        self.admission_var = tk.StringVar()
        self.learner_name_var = tk.StringVar()
        self.parent_name_var = tk.StringVar()
        self.parent_phone_var = tk.StringVar()
        self.class_level_var = tk.StringVar(value=CLASS_LEVELS[0])
        self.boarding_var = tk.StringVar(value="Day Scholar")
        self.transport_var = tk.StringVar(value="School Bus")
        self.learner_filter_var = tk.StringVar()
        self.learner_class_filter_var = tk.StringVar(value=ALL_CLASSES_FILTER)
        self.learner_boarding_filter_var = tk.StringVar(value=ALL_BOARDING_FILTER)
        self.learner_transport_filter_var = tk.StringVar(value=ALL_TRANSPORT_FILTER)

        self.reporting_search_var = tk.StringVar()
        self.reporting_class_filter_var = tk.StringVar(value=ALL_CLASSES_FILTER)
        self.reporting_boarding_filter_var = tk.StringVar(value=ALL_BOARDING_FILTER)
        self.reporting_transport_filter_var = tk.StringVar(value=ALL_TRANSPORT_FILTER)
        self.reporting_selected_var = tk.StringVar(value="Selected learner: None")
        self.reporting_date_var = tk.StringVar(value=now_date())
        self.reporting_time_var = tk.StringVar(value=now_time())
        self.reporting_transport_var = tk.StringVar(value="N/A")
        self.accompanied_mode_var = tk.StringVar(value="Registered Parent")
        self.accompanied_name_var = tk.StringVar()
        self.accompanied_phone_var = tk.StringVar()
        self.today_filter_var = tk.StringVar()

        today = now_date()
        self.report_from_var = tk.StringVar(value=today)
        self.report_to_var = tk.StringVar(value=today)
        self.report_type_var = tk.StringVar(value="boarders")
        self.report_format_var = tk.StringVar(value="PDF")
        self.report_orientation_var = tk.StringVar(
            value=self.db.get_setting("report_orientation", "Auto")
        )
        self.exam_portal_url_var = tk.StringVar(
            value=self.db.get_setting("exam_base_url", "http://127.0.0.1:5050")
        )

        self.setting_school_name_var = tk.StringVar(value=self.school_name)
        self.setting_report_dir_var = tk.StringVar(value=str(self.default_report_dir))
        self.setting_password_var = tk.StringVar(value=self.db.get_setting("app_password", ""))
        self.setting_logo_var = tk.StringVar(value=str(self.logo_path) if self.logo_path else "")
        self.setting_word_template_portrait_var = tk.StringVar(
            value=self.db.get_setting("word_template_portrait", "")
        )
        self.setting_word_template_landscape_var = tk.StringVar(
            value=self.db.get_setting("word_template_landscape", "")
        )
        self.setting_pdf_template_portrait_var = tk.StringVar(
            value=self.db.get_setting("pdf_template_portrait", "")
        )
        self.setting_pdf_template_landscape_var = tk.StringVar(
            value=self.db.get_setting("pdf_template_landscape", "")
        )
    def _build_menu(self) -> None:
        menubar = tk.Menu(self)

        file_menu = tk.Menu(menubar, tearoff=False)
        file_menu.add_command(label="Import Learners from Excel", command=self.import_learners_from_excel)
        file_menu.add_separator()
        file_menu.add_command(label="Backup Database", command=self.backup_database)
        file_menu.add_command(label="Restore Database", command=self.restore_database)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self._exit_app)
        menubar.add_cascade(label="File", menu=file_menu)

        reports_menu = tk.Menu(menubar, tearoff=False)
        reports_menu.add_command(label="Generate Current Report", command=self.generate_report)
        reports_menu.add_command(label="Open Last Report", command=self.open_last_report)
        reports_menu.add_command(label="Open Reports Folder", command=self.open_reports_folder)
        menubar.add_cascade(label="Reports", menu=reports_menu)

        exam_menu = tk.Menu(menubar, tearoff=False)
        exam_menu.add_command(label="Open Exam Marks App", command=self.launch_exam_app)
        exam_menu.add_command(label="Start Exam Portal", command=self.launch_exam_portal)
        exam_menu.add_command(label="Open Exam Portal in Browser", command=self.open_exam_portal_url)
        menubar.add_cascade(label="Exam", menu=exam_menu)

        help_menu = tk.Menu(menubar, tearoff=False)
        help_menu.add_command(label="User Manual", command=self.open_user_manual)
        help_menu.add_command(label="About", command=self.show_about)
        menubar.add_cascade(label="Help", menu=help_menu)
        self.config(menu=menubar)

    def _build_layout(self) -> None:
        container = ttk.Frame(self, padding=8)
        container.pack(fill=tk.BOTH, expand=True)

        header = ttk.Frame(container, padding=(4, 4, 4, 10))
        header.pack(fill=tk.X)
        self.logo_label = ttk.Label(header, text="Tumaini Logo")
        self.logo_label.pack(side=tk.LEFT, padx=(0, 10))
        title_text = f"{APP_TITLE}\nOffline Learner Reporting"
        self.header_title_label = ttk.Label(
            header,
            text=title_text,
            font=("Segoe UI", 12, "bold"),
            justify=tk.LEFT,
        )
        self.header_title_label.pack(side=tk.LEFT, anchor="w")

        self.notebook = ttk.Notebook(container)
        self.notebook.pack(fill=tk.BOTH, expand=True)

        self.learners_tab, self.learners_tab_content = self._create_scrollable_tab()
        self.reporting_tab, self.reporting_tab_content = self._create_scrollable_tab()
        self.reports_tab, self.reports_tab_content = self._create_scrollable_tab()
        self.exams_tab, self.exams_tab_content = self._create_scrollable_tab()
        self.backup_tab, self.backup_tab_content = self._create_scrollable_tab()

        self.notebook.add(self.learners_tab, text="Learners")
        self.notebook.add(self.reporting_tab, text="Reporting")
        self.notebook.add(self.reports_tab, text="Reports")
        self.notebook.add(self.exams_tab, text="Exams")
        self.notebook.add(self.backup_tab, text="Backup & Settings")

        self._build_learners_tab()
        self._build_reporting_tab()
        self._build_reports_tab()
        self._build_exams_tab()
        self._build_backup_tab()
        self._load_logo()

        status_bar = ttk.Label(self, textvariable=self.status_text, anchor="w")
        status_bar.pack(fill=tk.X, padx=8, pady=(0, 8))

    def _create_scrollable_tab(self) -> tuple[ttk.Frame, ttk.Frame]:
        tab_frame = ttk.Frame(self.notebook)
        canvas = tk.Canvas(tab_frame, highlightthickness=0, borderwidth=0)
        scrollbar = ttk.Scrollbar(tab_frame, orient=tk.VERTICAL, command=canvas.yview)
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        content = ttk.Frame(canvas, padding=10)
        content_window = canvas.create_window((0, 0), window=content, anchor="nw")

        def _on_content_configure(_event: tk.Event) -> None:
            canvas.configure(scrollregion=canvas.bbox("all"))

        def _on_canvas_configure(event: tk.Event) -> None:
            canvas.itemconfigure(content_window, width=event.width)

        content.bind("<Configure>", _on_content_configure)
        canvas.bind("<Configure>", _on_canvas_configure)
        canvas.bind("<Enter>", lambda _e: self._set_active_scroll_canvas(canvas))
        canvas.bind("<Leave>", lambda _e: self._clear_active_scroll_canvas(canvas))
        content.bind("<Enter>", lambda _e: self._set_active_scroll_canvas(canvas))
        return tab_frame, content

    def _set_active_scroll_canvas(self, canvas: tk.Canvas) -> None:
        self._active_mousewheel_canvas = canvas

    def _clear_active_scroll_canvas(self, canvas: tk.Canvas) -> None:
        if self._active_mousewheel_canvas is canvas:
            self._active_mousewheel_canvas = None

    def _on_global_mousewheel(self, event: tk.Event) -> None:
        canvas = getattr(self, "_active_mousewheel_canvas", None)
        if canvas is None or not canvas.winfo_exists():
            return
        scrollregion = canvas.bbox("all")
        if not scrollregion:
            return
        content_height = scrollregion[3] - scrollregion[1]
        if content_height <= canvas.winfo_height():
            return

        if hasattr(event, "num") and event.num == 4:
            delta_units = -1
        elif hasattr(event, "num") and event.num == 5:
            delta_units = 1
        else:
            delta = getattr(event, "delta", 0)
            if delta == 0:
                return
            delta_units = -1 * int(delta / 120) if sys.platform.startswith("win") else -1 * int(delta)
            if delta_units == 0:
                delta_units = -1 if delta > 0 else 1

        canvas.yview_scroll(delta_units, "units")
        return "break"

    def _build_learners_tab(self) -> None:
        form = ttk.LabelFrame(self.learners_tab_content, text="Learner Details", padding=10)
        form.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(form, text="Admission Number").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(form, textvariable=self.admission_var, width=24).grid(
            row=0, column=1, sticky="w", padx=5, pady=5
        )
        ttk.Button(form, text="Search", command=self.search_learner_by_admission).grid(
            row=0, column=2, sticky="w", padx=5, pady=5
        )

        ttk.Label(form, text="Learner Name").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(form, textvariable=self.learner_name_var, width=42).grid(
            row=1, column=1, columnspan=2, sticky="we", padx=5, pady=5
        )

        ttk.Label(form, text="Class Level").grid(row=2, column=0, sticky="w", padx=5, pady=5)
        ttk.Combobox(
            form,
            textvariable=self.class_level_var,
            values=CLASS_LEVELS,
            state="readonly",
            width=21,
        ).grid(row=2, column=1, sticky="w", padx=5, pady=5)

        ttk.Label(form, text="Parent Name").grid(row=3, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(form, textvariable=self.parent_name_var, width=42).grid(
            row=3, column=1, columnspan=2, sticky="we", padx=5, pady=5
        )

        ttk.Label(form, text="Parent Phone").grid(row=4, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(form, textvariable=self.parent_phone_var, width=24).grid(
            row=4, column=1, sticky="w", padx=5, pady=5
        )

        ttk.Label(form, text="Boarding Status").grid(row=5, column=0, sticky="w", padx=5, pady=5)
        boarding_combo = ttk.Combobox(
            form,
            textvariable=self.boarding_var,
            values=VALID_BOARDING,
            state="readonly",
            width=21,
        )
        boarding_combo.grid(row=5, column=1, sticky="w", padx=5, pady=5)
        boarding_combo.bind("<<ComboboxSelected>>", self._on_boarding_change)

        ttk.Label(form, text="Transport Mode").grid(row=6, column=0, sticky="w", padx=5, pady=5)
        self.transport_combo = ttk.Combobox(
            form,
            textvariable=self.transport_var,
            values=DAY_SCHOLAR_TRANSPORT,
            state="readonly",
            width=21,
        )
        self.transport_combo.grid(row=6, column=1, sticky="w", padx=5, pady=5)

        actions = ttk.Frame(form)
        actions.grid(row=7, column=0, columnspan=3, sticky="w", padx=5, pady=(8, 5))
        ttk.Button(actions, text="Save", command=self.save_learner).pack(side=tk.LEFT, padx=4)
        ttk.Button(actions, text="Clear", command=self.clear_learner_form).pack(side=tk.LEFT, padx=4)
        ttk.Button(actions, text="Update", command=self.update_learner).pack(side=tk.LEFT, padx=4)
        ttk.Button(actions, text="Delete", command=self.delete_learner).pack(side=tk.LEFT, padx=4)
        ttk.Button(actions, text="Import Excel List", command=self.import_learners_from_excel).pack(
            side=tk.LEFT,
            padx=4,
        )

        form.columnconfigure(1, weight=1)

        table_container = ttk.LabelFrame(self.learners_tab_content, text="Learners List", padding=10)
        table_container.pack(fill=tk.BOTH, expand=True)

        filter_row = ttk.Frame(table_container)
        filter_row.pack(fill=tk.X, pady=(0, 8))
        ttk.Label(filter_row, text="Filter").pack(side=tk.LEFT, padx=(0, 6))
        learner_filter_entry = ttk.Entry(filter_row, textvariable=self.learner_filter_var, width=32)
        learner_filter_entry.pack(side=tk.LEFT, padx=(0, 6))
        learner_filter_entry.bind("<KeyRelease>", lambda _e: self.refresh_learners_table())
        ttk.Label(filter_row, text="Class").pack(side=tk.LEFT, padx=(8, 4))
        learner_class_filter_combo = ttk.Combobox(
            filter_row,
            textvariable=self.learner_class_filter_var,
            values=(ALL_CLASSES_FILTER,) + CLASS_LEVELS,
            state="readonly",
            width=16,
        )
        learner_class_filter_combo.pack(side=tk.LEFT, padx=(0, 6))
        learner_class_filter_combo.bind("<<ComboboxSelected>>", lambda _e: self.refresh_learners_table())
        ttk.Label(filter_row, text="Boarding").pack(side=tk.LEFT, padx=(8, 4))
        learner_boarding_filter_combo = ttk.Combobox(
            filter_row,
            textvariable=self.learner_boarding_filter_var,
            values=(ALL_BOARDING_FILTER,) + VALID_BOARDING,
            state="readonly",
            width=14,
        )
        learner_boarding_filter_combo.pack(side=tk.LEFT, padx=(0, 6))
        learner_boarding_filter_combo.bind("<<ComboboxSelected>>", lambda _e: self.refresh_learners_table())
        ttk.Label(filter_row, text="Transport").pack(side=tk.LEFT, padx=(8, 4))
        learner_transport_filter_combo = ttk.Combobox(
            filter_row,
            textvariable=self.learner_transport_filter_var,
            values=(ALL_TRANSPORT_FILTER, "N/A") + DAY_SCHOLAR_TRANSPORT,
            state="readonly",
            width=14,
        )
        learner_transport_filter_combo.pack(side=tk.LEFT, padx=(0, 6))
        learner_transport_filter_combo.bind("<<ComboboxSelected>>", lambda _e: self.refresh_learners_table())
        ttk.Button(filter_row, text="Clear Filter", command=self.clear_learners_filter).pack(side=tk.LEFT)

        columns = (
            "no",
            "admission_no",
            "learner_name",
            "class_level",
            "parent_name",
            "parent_phone",
            "boarding_status",
            "transport_mode",
        )
        self.learners_tree = ttk.Treeview(
            table_container,
            columns=columns,
            show="headings",
            height=16,
        )
        headers = {
            "no": "No.",
            "admission_no": "Admission No.",
            "learner_name": "Learner Name",
            "class_level": "Class",
            "parent_name": "Parent Name",
            "parent_phone": "Parent Phone",
            "boarding_status": "Boarding Status",
            "transport_mode": "Transport Mode",
        }
        for col in columns:
            self.learners_tree.heading(col, text=headers[col])
            if col == "no":
                self.learners_tree.column(col, width=60, anchor="center")
            else:
                self.learners_tree.column(col, width=160, anchor="w")

        learner_scroll = ttk.Scrollbar(table_container, orient=tk.VERTICAL, command=self.learners_tree.yview)
        self.learners_tree.configure(yscrollcommand=learner_scroll.set)
        self.learners_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        learner_scroll.pack(side=tk.LEFT, fill=tk.Y)
        self.learners_tree.bind("<<TreeviewSelect>>", self._on_learner_tree_select)

        self._sync_transport_dropdown()
    def _build_reporting_tab(self) -> None:
        top = ttk.LabelFrame(self.reporting_tab_content, text="Search and Select Learner", padding=10)
        top.pack(fill=tk.BOTH, expand=True)

        search_row = ttk.Frame(top)
        search_row.pack(fill=tk.X, pady=(0, 8))
        ttk.Label(search_row, text="Search learner").pack(side=tk.LEFT, padx=(0, 8))
        search_entry = ttk.Entry(search_row, textvariable=self.reporting_search_var, width=34)
        search_entry.pack(side=tk.LEFT, padx=(0, 8))
        search_entry.bind("<KeyRelease>", lambda _e: self.refresh_reporting_learner_table())
        ttk.Button(search_row, text="Find", command=self.refresh_reporting_learner_table).pack(side=tk.LEFT)
        ttk.Button(search_row, text="Register New Learner", command=self.prompt_new_registration).pack(
            side=tk.LEFT, padx=(8, 0)
        )

        search_filter_row = ttk.Frame(top)
        search_filter_row.pack(fill=tk.X, pady=(0, 8))
        ttk.Label(search_filter_row, text="Class").pack(side=tk.LEFT, padx=(0, 4))
        report_class_filter_combo = ttk.Combobox(
            search_filter_row,
            textvariable=self.reporting_class_filter_var,
            values=(ALL_CLASSES_FILTER,) + CLASS_LEVELS,
            state="readonly",
            width=16,
        )
        report_class_filter_combo.pack(side=tk.LEFT, padx=(0, 8))
        report_class_filter_combo.bind("<<ComboboxSelected>>", lambda _e: self.refresh_reporting_learner_table())
        ttk.Label(search_filter_row, text="Boarding").pack(side=tk.LEFT, padx=(0, 4))
        report_boarding_filter_combo = ttk.Combobox(
            search_filter_row,
            textvariable=self.reporting_boarding_filter_var,
            values=(ALL_BOARDING_FILTER,) + VALID_BOARDING,
            state="readonly",
            width=14,
        )
        report_boarding_filter_combo.pack(side=tk.LEFT, padx=(0, 8))
        report_boarding_filter_combo.bind(
            "<<ComboboxSelected>>",
            lambda _e: self.refresh_reporting_learner_table(),
        )
        ttk.Label(search_filter_row, text="Transport").pack(side=tk.LEFT, padx=(0, 4))
        report_transport_filter_combo = ttk.Combobox(
            search_filter_row,
            textvariable=self.reporting_transport_filter_var,
            values=(ALL_TRANSPORT_FILTER, "N/A") + DAY_SCHOLAR_TRANSPORT,
            state="readonly",
            width=14,
        )
        report_transport_filter_combo.pack(side=tk.LEFT, padx=(0, 8))
        report_transport_filter_combo.bind(
            "<<ComboboxSelected>>",
            lambda _e: self.refresh_reporting_learner_table(),
        )
        ttk.Button(
            search_filter_row,
            text="Clear Filters",
            command=self.clear_reporting_filters,
        ).pack(side=tk.LEFT, padx=(4, 0))

        learner_columns = (
            "no",
            "admission_no",
            "learner_name",
            "class_level",
            "parent_name",
            "parent_phone",
            "boarding_status",
            "transport_mode",
        )
        self.reporting_learner_tree = ttk.Treeview(
            top,
            columns=learner_columns,
            show="headings",
            height=8,
        )
        report_headers = {
            "no": "No.",
            "admission_no": "Admission No.",
            "learner_name": "Learner Name",
            "class_level": "Class",
            "parent_name": "Parent Name",
            "parent_phone": "Parent Phone",
            "boarding_status": "Status",
            "transport_mode": "Transport",
        }
        for col in learner_columns:
            self.reporting_learner_tree.heading(col, text=report_headers[col])
            if col == "no":
                self.reporting_learner_tree.column(col, width=60, anchor="center")
            else:
                self.reporting_learner_tree.column(col, width=140, anchor="w")
        self.reporting_learner_tree.pack(fill=tk.BOTH, expand=True)
        self.reporting_learner_tree.bind("<<TreeviewSelect>>", self._on_reporting_learner_select)

        middle = ttk.LabelFrame(self.reporting_tab_content, text="Arrival Reporting", padding=10)
        middle.pack(fill=tk.X, pady=10)
        ttk.Label(middle, textvariable=self.reporting_selected_var, font=("Segoe UI", 10, "bold")).grid(
            row=0, column=0, columnspan=6, sticky="w", padx=5, pady=5
        )

        ttk.Label(middle, text="Date").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(middle, textvariable=self.reporting_date_var, state="readonly", width=16).grid(
            row=1, column=1, sticky="w", padx=5, pady=5
        )

        ttk.Label(middle, text="Reporting Time (HH:MM)").grid(row=1, column=2, sticky="w", padx=5, pady=5)
        ttk.Entry(middle, textvariable=self.reporting_time_var, width=16).grid(
            row=1, column=3, sticky="w", padx=5, pady=5
        )

        ttk.Label(middle, text="Arrival Transport").grid(row=1, column=4, sticky="w", padx=5, pady=5)
        self.reporting_transport_combo = ttk.Combobox(
            middle,
            textvariable=self.reporting_transport_var,
            values=("N/A",) + DAY_SCHOLAR_TRANSPORT,
            state="readonly",
            width=16,
        )
        self.reporting_transport_combo.grid(row=1, column=5, sticky="w", padx=5, pady=5)

        ttk.Label(middle, text="Accompanied By").grid(row=2, column=0, sticky="w", padx=5, pady=5)
        self.accompanied_mode_combo = ttk.Combobox(
            middle,
            textvariable=self.accompanied_mode_var,
            values=("Registered Parent", "Other Person"),
            state="readonly",
            width=16,
        )
        self.accompanied_mode_combo.grid(row=2, column=1, sticky="w", padx=5, pady=5)
        self.accompanied_mode_combo.bind("<<ComboboxSelected>>", lambda _e: self._sync_accompanied_fields())

        ttk.Label(middle, text="Accompanied Name").grid(row=2, column=2, sticky="w", padx=5, pady=5)
        self.accompanied_name_entry = ttk.Entry(middle, textvariable=self.accompanied_name_var, width=40)
        self.accompanied_name_entry.grid(row=2, column=3, columnspan=3, sticky="we", padx=5, pady=5)

        ttk.Label(middle, text="Accompanied Phone").grid(row=3, column=0, sticky="w", padx=5, pady=5)
        self.accompanied_phone_entry = ttk.Entry(middle, textvariable=self.accompanied_phone_var, width=18)
        self.accompanied_phone_entry.grid(row=3, column=1, sticky="w", padx=5, pady=5)

        ttk.Button(middle, text="Use Current Time", command=self.reset_reporting_time).grid(
            row=4, column=0, sticky="w", padx=5, pady=(8, 5)
        )
        ttk.Button(middle, text="Reported", command=self.report_arrival).grid(
            row=4, column=1, sticky="w", padx=5, pady=(8, 5)
        )

        middle.columnconfigure(1, weight=1)

        bottom = ttk.LabelFrame(self.reporting_tab_content, text="Today's Reported Learners", padding=10)
        bottom.pack(fill=tk.BOTH, expand=True)

        filter_row = ttk.Frame(bottom)
        filter_row.pack(fill=tk.X, pady=(0, 8))
        ttk.Label(filter_row, text="Filter").pack(side=tk.LEFT, padx=(0, 6))
        filter_entry = ttk.Entry(filter_row, textvariable=self.today_filter_var, width=32)
        filter_entry.pack(side=tk.LEFT, padx=(0, 8))
        filter_entry.bind("<KeyRelease>", lambda _e: self.refresh_today_reporting_table())
        ttk.Button(filter_row, text="Refresh", command=self.refresh_today_reporting_table).pack(side=tk.LEFT)
        ttk.Button(
            filter_row,
            text="Delete Selected",
            command=self.delete_selected_reported_entry,
        ).pack(side=tk.LEFT, padx=(8, 0))

        today_columns = (
            "no",
            "report_time",
            "learner_name",
            "admission_no",
            "arrival_transport_mode",
            "accompanied_by_name",
            "accompanied_by_phone",
        )
        self.today_tree = ttk.Treeview(bottom, columns=today_columns, show="headings", height=10)
        today_headers = {
            "no": "No.",
            "report_time": "Time",
            "learner_name": "Learner Name",
            "admission_no": "Admission No.",
            "arrival_transport_mode": "Transport",
            "accompanied_by_name": "Accompanied By",
            "accompanied_by_phone": "Accompanied Phone",
        }
        for col in today_columns:
            self.today_tree.heading(col, text=today_headers[col])
            if col == "no":
                self.today_tree.column(col, width=60, anchor="center")
            else:
                self.today_tree.column(col, width=170, anchor="w")
        self.today_tree.pack(fill=tk.BOTH, expand=True)

    def _build_reports_tab(self) -> None:
        controls = ttk.LabelFrame(self.reports_tab_content, text="Report Parameters", padding=10)
        controls.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(controls, text="From Date (YYYY-MM-DD)").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(controls, textvariable=self.report_from_var, width=18).grid(
            row=0, column=1, sticky="w", padx=5, pady=5
        )
        ttk.Label(controls, text="To Date (YYYY-MM-DD)").grid(row=0, column=2, sticky="w", padx=5, pady=5)
        ttk.Entry(controls, textvariable=self.report_to_var, width=18).grid(
            row=0, column=3, sticky="w", padx=5, pady=5
        )

        report_type_box = ttk.LabelFrame(controls, text="Report Type", padding=8)
        report_type_box.grid(row=1, column=0, columnspan=4, sticky="we", padx=5, pady=6)
        ttk.Radiobutton(
            report_type_box, text="Boarders List", variable=self.report_type_var, value="boarders"
        ).pack(anchor="w", padx=4, pady=2)
        ttk.Radiobutton(
            report_type_box, text="School Bus Users", variable=self.report_type_var, value="school_bus"
        ).pack(anchor="w", padx=4, pady=2)
        ttk.Radiobutton(
            report_type_box,
            text="All Learners with Parent Details",
            variable=self.report_type_var,
            value="all_learners",
        ).pack(anchor="w", padx=4, pady=2)
        ttk.Radiobutton(
            report_type_box,
            text="Parents with Multiple Learners",
            variable=self.report_type_var,
            value="multi_parents",
        ).pack(anchor="w", padx=4, pady=2)
        ttk.Radiobutton(
            report_type_box,
            text="Reported Learners (Date Range)",
            variable=self.report_type_var,
            value="reported_learners",
        ).pack(anchor="w", padx=4, pady=2)

        ttk.Label(controls, text="Format").grid(row=2, column=0, sticky="w", padx=5, pady=(6, 5))
        ttk.Combobox(
            controls,
            textvariable=self.report_format_var,
            values=("PDF", "Word"),
            state="readonly",
            width=16,
        ).grid(row=2, column=1, sticky="w", padx=5, pady=(6, 5))

        ttk.Label(controls, text="Orientation").grid(row=2, column=2, sticky="w", padx=5, pady=(6, 5))
        ttk.Combobox(
            controls,
            textvariable=self.report_orientation_var,
            values=REPORT_ORIENTATIONS,
            state="readonly",
            width=16,
        ).grid(row=2, column=3, sticky="w", padx=5, pady=(6, 5))

        ttk.Button(controls, text="Generate Report", command=self.generate_report).grid(
            row=2, column=4, sticky="w", padx=5, pady=(6, 5)
        )
        ttk.Button(controls, text="Open Last Report", command=self.open_last_report).grid(
            row=2, column=5, sticky="w", padx=5, pady=(6, 5)
        )
        ttk.Button(controls, text="Open Reports Folder", command=self.open_reports_folder).grid(
            row=2, column=6, sticky="w", padx=5, pady=(6, 5)
        )

        dependency_note = "PDF support: OK" if REPORTLAB_AVAILABLE else "PDF support: Missing reportlab"
        dependency_note += " | Word support: OK" if DOCX_AVAILABLE else " | Word support: Missing python-docx"
        dependency_note += " | PDF templates: PNG/JPG background | Word templates: DOCX"
        ttk.Label(controls, text=dependency_note).grid(
            row=3, column=0, columnspan=7, sticky="w", padx=5, pady=(4, 2)
        )
        controls.columnconfigure(6, weight=1)

        preview_box = ttk.LabelFrame(self.reports_tab_content, text="Report Preview", padding=10)
        preview_box.pack(fill=tk.BOTH, expand=True)
        self.report_preview = ScrolledText(preview_box, wrap=tk.WORD, height=20)
        self.report_preview.pack(fill=tk.BOTH, expand=True)
        self.report_preview.insert("1.0", "Generated report data will appear here.")
        self.report_preview.configure(state=tk.DISABLED)

    def _build_exams_tab(self) -> None:
        launcher_box = ttk.LabelFrame(self.exams_tab_content, text="Exam System Launcher", padding=10)
        launcher_box.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(
            launcher_box,
            text=(
                "Use this section to open the separate Exam Marks app and the online teacher portal.\n"
                "The exam system uses the same learner list and class levels already in this LMS."
            ),
            justify=tk.LEFT,
        ).grid(row=0, column=0, columnspan=3, sticky="w", padx=5, pady=(0, 10))

        ttk.Button(
            launcher_box,
            text="Open Exam Marks Desktop App",
            command=self.launch_exam_app,
        ).grid(row=1, column=0, sticky="w", padx=5, pady=5)

        ttk.Button(
            launcher_box,
            text="Start Online Exam Portal",
            command=self.launch_exam_portal,
        ).grid(row=1, column=1, sticky="w", padx=5, pady=5)

        ttk.Button(
            launcher_box,
            text="Open Portal in Browser",
            command=self.open_exam_portal_url,
        ).grid(row=1, column=2, sticky="w", padx=5, pady=5)

        portal_box = ttk.LabelFrame(self.exams_tab_content, text="Portal URL", padding=10)
        portal_box.pack(fill=tk.X)

        ttk.Label(portal_box, text="Exam Portal Base URL").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(portal_box, textvariable=self.exam_portal_url_var, width=46).grid(
            row=0, column=1, sticky="we", padx=5, pady=5
        )
        ttk.Button(portal_box, text="Save URL", command=self.save_exam_portal_url).grid(
            row=0, column=2, sticky="w", padx=5, pady=5
        )
        ttk.Label(
            portal_box,
            text="Teachers receive links generated from the Exam app: <BaseURL>/fill/<token>",
        ).grid(row=1, column=0, columnspan=3, sticky="w", padx=5, pady=(0, 5))
        portal_box.columnconfigure(1, weight=1)

    def _build_backup_tab(self) -> None:
        backup_box = ttk.LabelFrame(self.backup_tab_content, text="Database Backup and Restore", padding=10)
        backup_box.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(backup_box, text=f"Active Database: {DB_PATH}").pack(anchor="w", pady=(0, 8))
        ttk.Button(backup_box, text="Backup Database", command=self.backup_database).pack(
            side=tk.LEFT, padx=(0, 6)
        )
        ttk.Button(backup_box, text="Restore from Backup", command=self.restore_database).pack(
            side=tk.LEFT, padx=(0, 6)
        )

        settings_box = ttk.LabelFrame(self.backup_tab_content, text="Settings", padding=10)
        settings_box.pack(fill=tk.X)

        ttk.Label(settings_box, text="School Name").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(settings_box, textvariable=self.setting_school_name_var, width=42).grid(
            row=0, column=1, columnspan=3, sticky="we", padx=5, pady=5
        )

        ttk.Label(settings_box, text="Default Report Folder").grid(
            row=1, column=0, sticky="w", padx=5, pady=5
        )
        ttk.Entry(settings_box, textvariable=self.setting_report_dir_var, width=42).grid(
            row=1, column=1, columnspan=2, sticky="we", padx=5, pady=5
        )
        ttk.Button(settings_box, text="Browse", command=self.browse_report_folder).grid(
            row=1, column=3, sticky="w", padx=5, pady=5
        )

        ttk.Label(settings_box, text="Logo Image Path").grid(
            row=2, column=0, sticky="w", padx=5, pady=5
        )
        ttk.Entry(settings_box, textvariable=self.setting_logo_var, width=42).grid(
            row=2, column=1, columnspan=2, sticky="we", padx=5, pady=5
        )
        ttk.Button(settings_box, text="Browse", command=self.browse_logo_file).grid(
            row=2, column=3, sticky="w", padx=5, pady=5
        )

        ttk.Label(settings_box, text="Optional App Password").grid(
            row=7, column=0, sticky="w", padx=5, pady=5
        )
        ttk.Entry(settings_box, textvariable=self.setting_password_var, show="*", width=28).grid(
            row=7, column=1, sticky="w", padx=5, pady=5
        )
        ttk.Label(settings_box, text="Leave empty to disable login").grid(
            row=7, column=2, columnspan=2, sticky="w", padx=5, pady=5
        )

        ttk.Label(settings_box, text="Word Template (Portrait)").grid(
            row=3, column=0, sticky="w", padx=5, pady=5
        )
        ttk.Entry(settings_box, textvariable=self.setting_word_template_portrait_var, width=42).grid(
            row=3, column=1, columnspan=2, sticky="we", padx=5, pady=5
        )
        ttk.Button(
            settings_box,
            text="Browse",
            command=lambda: self.browse_template_file(
                self.setting_word_template_portrait_var,
                "Select Word Portrait Template",
                [("Word Document", "*.docx"), ("All Files", "*.*")],
            ),
        ).grid(row=3, column=3, sticky="w", padx=5, pady=5)

        ttk.Label(settings_box, text="Word Template (Landscape)").grid(
            row=4, column=0, sticky="w", padx=5, pady=5
        )
        ttk.Entry(settings_box, textvariable=self.setting_word_template_landscape_var, width=42).grid(
            row=4, column=1, columnspan=2, sticky="we", padx=5, pady=5
        )
        ttk.Button(
            settings_box,
            text="Browse",
            command=lambda: self.browse_template_file(
                self.setting_word_template_landscape_var,
                "Select Word Landscape Template",
                [("Word Document", "*.docx"), ("All Files", "*.*")],
            ),
        ).grid(row=4, column=3, sticky="w", padx=5, pady=5)

        ttk.Label(settings_box, text="PDF Template (Portrait)").grid(
            row=5, column=0, sticky="w", padx=5, pady=5
        )
        ttk.Entry(settings_box, textvariable=self.setting_pdf_template_portrait_var, width=42).grid(
            row=5, column=1, columnspan=2, sticky="we", padx=5, pady=5
        )
        ttk.Button(
            settings_box,
            text="Browse",
            command=lambda: self.browse_template_file(
                self.setting_pdf_template_portrait_var,
                "Select PDF Portrait Template (Image)",
                [("Image Files", "*.png;*.jpg;*.jpeg"), ("All Files", "*.*")],
            ),
        ).grid(row=5, column=3, sticky="w", padx=5, pady=5)

        ttk.Label(settings_box, text="PDF Template (Landscape)").grid(
            row=6, column=0, sticky="w", padx=5, pady=5
        )
        ttk.Entry(settings_box, textvariable=self.setting_pdf_template_landscape_var, width=42).grid(
            row=6, column=1, columnspan=2, sticky="we", padx=5, pady=5
        )
        ttk.Button(
            settings_box,
            text="Browse",
            command=lambda: self.browse_template_file(
                self.setting_pdf_template_landscape_var,
                "Select PDF Landscape Template (Image)",
                [("Image Files", "*.png;*.jpg;*.jpeg"), ("All Files", "*.*")],
            ),
        ).grid(row=6, column=3, sticky="w", padx=5, pady=5)

        ttk.Label(
            settings_box,
            text="Auto orientation switches to landscape when table is wide.",
        ).grid(
            row=8, column=0, columnspan=4, sticky="w", padx=5, pady=(2, 2)
        )

        ttk.Button(settings_box, text="Save Settings", command=self.save_settings).grid(
            row=9, column=0, sticky="w", padx=5, pady=(8, 4)
        )
        settings_box.columnconfigure(1, weight=1)
    def _check_optional_login(self) -> bool:
        password = self.db.get_setting("app_password", "").strip()
        if not password:
            return True
        for _attempt in range(3):
            entered = simpledialog.askstring(
                "Login Required",
                "Enter system password:",
                show="*",
                parent=self,
            )
            if entered is None:
                return False
            if entered == password:
                return True
            messagebox.showerror("Login Failed", "Incorrect password.")
        messagebox.showerror("Access Denied", "Too many failed attempts.")
        return False

    def _load_initial_data(self) -> None:
        self.refresh_learners_table()
        self.refresh_reporting_learner_table()
        self.refresh_today_reporting_table()
        self._sync_accompanied_fields()
        self._sync_reporting_transport()

    def _load_logo(self) -> None:
        logo_text = "Tumaini Academy"
        logo_value = self.setting_logo_var.get().strip() if hasattr(self, "setting_logo_var") else ""
        logo_path = Path(logo_value) if logo_value else self.logo_path
        if logo_path and logo_path.exists():
            try:
                source = tk.PhotoImage(file=str(logo_path))
                scale = max(source.width() // 150, 1)
                self.logo_image = source.subsample(scale, scale)
                self.logo_label.configure(image=self.logo_image, text="")
                self.logo_path = logo_path
                return
            except tk.TclError:
                pass
        self.logo_image = None
        self.logo_label.configure(image="", text=logo_text)

    def _sync_reporting_transport(self) -> None:
        status = (self.selected_reporting_status or "").strip()
        if not status:
            self.reporting_transport_combo.configure(values=("N/A",) + DAY_SCHOLAR_TRANSPORT, state="disabled")
            self.reporting_transport_var.set("N/A")
            return
        if status == "Boarder":
            self.reporting_transport_combo.configure(values=("N/A",), state="disabled")
            self.reporting_transport_var.set("N/A")
            return
        self.reporting_transport_combo.configure(
            values=DAY_SCHOLAR_TRANSPORT,
            state="readonly",
        )
        selected = self.reporting_transport_var.get().strip()
        if selected not in DAY_SCHOLAR_TRANSPORT:
            fallback = self.selected_reporting_transport or "School Bus"
            if fallback not in DAY_SCHOLAR_TRANSPORT:
                fallback = "School Bus"
            self.reporting_transport_var.set(fallback)

    def _sync_accompanied_fields(self) -> None:
        mode = self.accompanied_mode_var.get()
        if mode == "Registered Parent":
            self.accompanied_name_var.set(self.selected_reporting_parent_name)
            self.accompanied_phone_var.set(self.selected_reporting_parent_phone)
            self.accompanied_name_entry.configure(state="disabled")
            self.accompanied_phone_entry.configure(state="disabled")
            return
        if (
            self.accompanied_name_var.get() == self.selected_reporting_parent_name
            and self.accompanied_phone_var.get() == self.selected_reporting_parent_phone
        ):
            self.accompanied_name_var.set("")
            self.accompanied_phone_var.set("")
        self.accompanied_name_entry.configure(state="normal")
        self.accompanied_phone_entry.configure(state="normal")

    def prompt_new_registration(self) -> None:
        search_value = self.reporting_search_var.get().strip()
        self.notebook.select(self.learners_tab)
        if search_value:
            self.admission_var.set(normalize_admission_no(search_value))
        self._set_status("Open Learners tab to register a new learner.")

    def import_learners_from_excel(self) -> None:
        initial_dir = str(DEFAULT_IMPORT_XLSX_PATH.parent if DEFAULT_IMPORT_XLSX_PATH.exists() else BASE_DIR)
        initial_file = DEFAULT_IMPORT_XLSX_PATH.name if DEFAULT_IMPORT_XLSX_PATH.exists() else ""
        file_path = filedialog.askopenfilename(
            title="Select Learners Excel File",
            initialdir=initial_dir,
            initialfile=initial_file,
            filetypes=[("Excel Workbook", "*.xlsx"), ("All Files", "*.*")],
        )
        if not file_path:
            return

        try:
            rows = read_learners_from_xlsx(Path(file_path))
        except Exception as error:
            messagebox.showerror("Import Failed", str(error))
            return

        inserted = 0
        skipped_existing = 0
        skipped_invalid = 0
        seen_in_file: set[str] = set()

        for row in rows:
            admission_no = normalize_admission_no(row.get("admission_no", ""))
            learner_name = row.get("learner_name", "").strip().title()
            class_level = normalize_class_level(row.get("class_level", ""))
            parent_name = row.get("parent_name", "").strip().title() or "Unknown Parent"
            parent_phone = sanitize_phone_for_storage(row.get("parent_phone", ""))

            if not admission_no or not learner_name or not class_level:
                skipped_invalid += 1
                continue
            if admission_no in seen_in_file:
                skipped_invalid += 1
                continue
            seen_in_file.add(admission_no)

            ok, _ = self.db.add_learner(
                {
                    "admission_no": admission_no,
                    "learner_name": learner_name,
                    "class_level": class_level,
                    "parent_name": parent_name,
                    "parent_phone": parent_phone,
                    "boarding_status": "Day Scholar",
                    "transport_mode": "School Bus",
                }
            )
            if ok:
                inserted += 1
            else:
                skipped_existing += 1

        self.refresh_learners_table()
        self.refresh_reporting_learner_table()
        summary = (
            f"Import completed.\n\nAdded: {inserted}\n"
            f"Skipped existing: {skipped_existing}\nSkipped invalid/duplicate rows: {skipped_invalid}"
        )
        self._set_status(summary.replace("\n", " | "))
        messagebox.showinfo("Excel Import", summary)

    def _set_status(self, text: str) -> None:
        self.status_text.set(text)

    def _clear_tree(self, tree: ttk.Treeview) -> None:
        tree.delete(*tree.get_children())

    def _sync_transport_dropdown(self) -> None:
        if self.boarding_var.get() == "Boarder":
            self.transport_combo.configure(values=("N/A",))
            self.transport_var.set("N/A")
            self.transport_combo.configure(state="disabled")
        else:
            self.transport_combo.configure(state="readonly", values=DAY_SCHOLAR_TRANSPORT)
            if self.transport_var.get() not in DAY_SCHOLAR_TRANSPORT:
                self.transport_var.set("School Bus")

    def _on_boarding_change(self, _event: tk.Event | None = None) -> None:
        self._sync_transport_dropdown()

    def _on_learner_tree_select(self, _event: tk.Event | None = None) -> None:
        selected = self.learners_tree.selection()
        if not selected:
            return
        row = self.learners_tree.item(selected[0], "values")
        if not row:
            return
        self.admission_var.set(str(row[1]))
        self.learner_name_var.set(str(row[2]))
        self.class_level_var.set(str(row[3]))
        self.parent_name_var.set(str(row[4]))
        self.parent_phone_var.set(str(row[5]))
        self.boarding_var.set(str(row[6]))
        self.transport_var.set(str(row[7]))
        self._sync_transport_dropdown()

    def _on_reporting_learner_select(self, _event: tk.Event | None = None) -> None:
        selected = self.reporting_learner_tree.selection()
        if not selected:
            return
        row = self.reporting_learner_tree.item(selected[0], "values")
        if not row:
            return
        self.selected_reporting_adm = str(row[1])
        self.selected_reporting_parent_name = str(row[4])
        self.selected_reporting_parent_phone = str(row[5])
        self.selected_reporting_status = str(row[6])
        self.selected_reporting_transport = str(row[7])
        self.reporting_selected_var.set(
            f"Selected learner: {row[2]} ({row[1]}) | Class: {row[3]} | Status: {row[6]}"
        )
        self.accompanied_mode_var.set("Registered Parent")
        self._sync_reporting_transport()
        self._sync_accompanied_fields()

    def clear_learners_filter(self) -> None:
        self.learner_filter_var.set("")
        self.learner_class_filter_var.set(ALL_CLASSES_FILTER)
        self.learner_boarding_filter_var.set(ALL_BOARDING_FILTER)
        self.learner_transport_filter_var.set(ALL_TRANSPORT_FILTER)
        self.refresh_learners_table()

    def clear_reporting_filters(self) -> None:
        self.reporting_search_var.set("")
        self.reporting_class_filter_var.set(ALL_CLASSES_FILTER)
        self.reporting_boarding_filter_var.set(ALL_BOARDING_FILTER)
        self.reporting_transport_filter_var.set(ALL_TRANSPORT_FILTER)
        self.refresh_reporting_learner_table()

    def refresh_learners_table(self) -> None:
        rows = self.db.search_learners(self.learner_filter_var.get().strip())
        self._clear_tree(self.learners_tree)
        class_filter = self.learner_class_filter_var.get()
        boarding_filter = self.learner_boarding_filter_var.get()
        transport_filter = self.learner_transport_filter_var.get()
        display_no = 1
        for row in rows:
            if class_filter != ALL_CLASSES_FILTER and row["class_level"] != class_filter:
                continue
            if boarding_filter != ALL_BOARDING_FILTER and row["boarding_status"] != boarding_filter:
                continue
            if transport_filter != ALL_TRANSPORT_FILTER and row["transport_mode"] != transport_filter:
                continue
            self.learners_tree.insert(
                "",
                tk.END,
                values=(
                    display_no,
                    row["admission_no"],
                    row["learner_name"],
                    row["class_level"],
                    row["parent_name"],
                    row["parent_phone"],
                    row["boarding_status"],
                    row["transport_mode"],
                ),
            )
            display_no += 1

    def refresh_reporting_learner_table(self) -> None:
        rows = self.db.search_learners(self.reporting_search_var.get().strip())
        self._clear_tree(self.reporting_learner_tree)
        class_filter = self.reporting_class_filter_var.get()
        boarding_filter = self.reporting_boarding_filter_var.get()
        transport_filter = self.reporting_transport_filter_var.get()
        display_no = 1
        for row in rows:
            if class_filter != ALL_CLASSES_FILTER and row["class_level"] != class_filter:
                continue
            if boarding_filter != ALL_BOARDING_FILTER and row["boarding_status"] != boarding_filter:
                continue
            if transport_filter != ALL_TRANSPORT_FILTER and row["transport_mode"] != transport_filter:
                continue
            self.reporting_learner_tree.insert(
                "",
                tk.END,
                values=(
                    display_no,
                    row["admission_no"],
                    row["learner_name"],
                    row["class_level"],
                    row["parent_name"],
                    row["parent_phone"],
                    row["boarding_status"],
                    row["transport_mode"],
                ),
            )
            display_no += 1

    def refresh_today_reporting_table(self) -> None:
        term = self.today_filter_var.get().strip()
        rows = self.db.get_reporting_for_date(self.reporting_date_var.get(), term)
        self._clear_tree(self.today_tree)
        for display_no, row in enumerate(rows, start=1):
            self.today_tree.insert(
                "",
                tk.END,
                iid=str(row["id"]),
                values=(
                    display_no,
                    row["report_time"],
                    row["learner_name"],
                    row["admission_no"],
                    row["arrival_transport_mode"],
                    row["accompanied_by_name"],
                    row["accompanied_by_phone"],
                ),
            )

    def delete_selected_reported_entry(self) -> None:
        selected = self.today_tree.selection()
        if not selected:
            messagebox.showwarning("Delete Reporting", "Select a reported learner entry to delete.")
            return

        item_id = selected[0]
        item_values = self.today_tree.item(item_id, "values")
        learner_label = item_values[2] if len(item_values) > 2 else "selected learner"

        confirm = messagebox.askyesno(
            "Confirm Delete",
            f"Delete reported entry for {learner_label}?",
        )
        if not confirm:
            return

        try:
            record_id = int(item_id)
        except ValueError:
            messagebox.showerror("Delete Reporting", "Invalid reporting record selected.")
            return

        ok, message = self.db.delete_reporting_record(record_id)
        if ok:
            self.refresh_today_reporting_table()
            self._set_status(message)
            messagebox.showinfo("Delete Reporting", message)
        else:
            messagebox.showerror("Delete Reporting", message)

    def _validate_learner(self) -> dict[str, str] | None:
        admission_no = normalize_admission_no(self.admission_var.get())
        learner_name = self.learner_name_var.get().strip()
        class_level = self.class_level_var.get().strip()
        parent_name = self.parent_name_var.get().strip()
        parent_phone = normalize_phone(self.parent_phone_var.get().strip())
        boarding = self.boarding_var.get().strip()
        transport = self.transport_var.get().strip()

        if not admission_no:
            messagebox.showerror("Validation Error", "Admission number is required.")
            return None
        if not learner_name:
            messagebox.showerror("Validation Error", "Learner name is required.")
            return None
        if class_level not in CLASS_LEVELS:
            messagebox.showerror("Validation Error", "Select a valid class level.")
            return None
        if not parent_name:
            messagebox.showerror("Validation Error", "Parent name is required.")
            return None
        if not PHONE_REGEX.match(parent_phone):
            messagebox.showerror(
                "Validation Error",
                "Invalid parent phone format. Use 10-15 digits, optional leading +.",
            )
            return None
        if boarding not in VALID_BOARDING:
            messagebox.showerror("Validation Error", "Select a valid boarding status.")
            return None
        if boarding == "Boarder":
            transport = "N/A"
        elif transport not in DAY_SCHOLAR_TRANSPORT:
            messagebox.showerror("Validation Error", "Select a valid transport mode for day scholars.")
            return None

        self.admission_var.set(admission_no)
        self.parent_phone_var.set(parent_phone)

        return {
            "admission_no": admission_no,
            "learner_name": learner_name,
            "class_level": class_level,
            "parent_name": parent_name,
            "parent_phone": parent_phone,
            "boarding_status": boarding,
            "transport_mode": transport,
        }

    def clear_learner_form(self) -> None:
        self.admission_var.set("")
        self.learner_name_var.set("")
        self.class_level_var.set(CLASS_LEVELS[0])
        self.parent_name_var.set("")
        self.parent_phone_var.set("")
        self.boarding_var.set("Day Scholar")
        self.transport_var.set("School Bus")
        self._sync_transport_dropdown()

    def save_learner(self) -> None:
        data = self._validate_learner()
        if not data:
            return
        ok, message = self.db.add_learner(data)
        if ok:
            messagebox.showinfo("Success", message)
            self._set_status(message)
            self.refresh_learners_table()
            self.refresh_reporting_learner_table()
            self.clear_learner_form()
        else:
            messagebox.showerror("Save Failed", message)
            self._set_status(message)

    def search_learner_by_admission(self) -> None:
        admission_no = normalize_admission_no(self.admission_var.get())
        if not admission_no:
            messagebox.showwarning("Search", "Enter admission number to search.")
            return
        row = self.db.get_learner(admission_no)
        if not row:
            messagebox.showinfo("Search", f"No learner found with admission number '{admission_no}'.")
            return
        self.admission_var.set(row["admission_no"])
        self.learner_name_var.set(row["learner_name"])
        self.class_level_var.set(row["class_level"])
        self.parent_name_var.set(row["parent_name"])
        self.parent_phone_var.set(row["parent_phone"])
        self.boarding_var.set(row["boarding_status"])
        self.transport_var.set(row["transport_mode"])
        self._sync_transport_dropdown()
        self._set_status(f"Learner '{admission_no}' loaded.")

    def update_learner(self) -> None:
        data = self._validate_learner()
        if not data:
            return
        ok, message = self.db.update_learner(data)
        if ok:
            messagebox.showinfo("Success", message)
            self._set_status(message)
            self.refresh_learners_table()
            self.refresh_reporting_learner_table()
        else:
            messagebox.showerror("Update Failed", message)
            self._set_status(message)

    def delete_learner(self) -> None:
        admission_no = self.admission_var.get().strip()
        if not admission_no:
            messagebox.showwarning("Delete", "Enter or select a learner first.")
            return
        confirm = messagebox.askyesno(
            "Confirm Delete",
            f"Delete learner '{admission_no}' and related reporting records?",
        )
        if not confirm:
            return
        ok, message = self.db.delete_learner(admission_no)
        if ok:
            messagebox.showinfo("Deleted", message)
            self._set_status(message)
            self.refresh_learners_table()
            self.refresh_reporting_learner_table()
            self.refresh_today_reporting_table()
            self.clear_learner_form()
        else:
            messagebox.showerror("Delete Failed", message)
            self._set_status(message)
    def reset_reporting_time(self) -> None:
        self.reporting_time_var.set(now_time())

    def report_arrival(self) -> None:
        if not self.selected_reporting_adm:
            query = self.reporting_search_var.get().strip()
            if query:
                register = messagebox.askyesno(
                    "Learner Not Found",
                    f"No learner selected for '{query}'.\nDo you want to register this learner now?",
                )
                if register:
                    self.prompt_new_registration()
            else:
                messagebox.showwarning("Reporting", "Search and select a learner first.")
            return
        row = self.db.get_learner(self.selected_reporting_adm)
        if not row:
            messagebox.showerror("Reporting", "Selected learner no longer exists.")
            self.refresh_reporting_learner_table()
            return

        report_date = self.reporting_date_var.get().strip()
        report_time = self.reporting_time_var.get().strip()

        try:
            parse_date(report_date)
        except ValueError:
            messagebox.showerror("Validation Error", "Date must be in YYYY-MM-DD format.")
            return

        if not is_valid_time(report_time):
            messagebox.showerror("Validation Error", "Reporting time must be in HH:MM format.")
            return

        if row["boarding_status"] == "Boarder":
            arrival_transport = "N/A"
        else:
            arrival_transport = self.reporting_transport_var.get().strip()
            if arrival_transport not in DAY_SCHOLAR_TRANSPORT:
                messagebox.showerror(
                    "Validation Error",
                    "Select transport mode for day scholar reporting.",
                )
                return

        accompanied_source = self.accompanied_mode_var.get().strip()
        accompanied_name = self.accompanied_name_var.get().strip()
        accompanied_phone = normalize_phone(self.accompanied_phone_var.get().strip())
        if accompanied_source not in ("Registered Parent", "Other Person"):
            accompanied_source = "Registered Parent"

        if accompanied_source == "Registered Parent":
            accompanied_name = row["parent_name"]
            accompanied_phone = row["parent_phone"]
        else:
            if not accompanied_name:
                messagebox.showerror("Validation Error", "Enter the name of the person accompanying learner.")
                return
            if accompanied_phone and not PHONE_REGEX.match(accompanied_phone):
                messagebox.showerror(
                    "Validation Error",
                    "Accompanied phone format is invalid. Use 10-15 digits, optional leading +.",
                )
                return

        company = accompanied_name or accompanied_source

        ok, message = self.db.add_reporting_record(
            admission_no=row["admission_no"],
            learner_name=row["learner_name"],
            report_date=report_date,
            report_time=report_time,
            arrival_transport_mode=arrival_transport,
            accompanied_source=accompanied_source,
            accompanied_by_name=accompanied_name,
            accompanied_by_phone=accompanied_phone,
            company=company,
        )
        if not ok:
            messagebox.showwarning("Reporting", message)
            self._set_status(message)
            return
        self.accompanied_mode_var.set("Registered Parent")
        self._sync_accompanied_fields()
        self.reset_reporting_time()
        self.refresh_today_reporting_table()
        self._set_status(f"Arrival reported for {row['learner_name']} ({row['admission_no']}).")

    def _prepare_report_data(self) -> tuple[str, list[str], list[list[str]], str, str]:
        report_kind = self.report_type_var.get()
        from_date = self.report_from_var.get().strip()
        to_date = self.report_to_var.get().strip()

        if report_kind == "reported_learners" and (not from_date or not to_date):
            raise ValueError("From Date and To Date are required for this report type.")

        try:
            if from_date:
                parse_date(from_date)
            if to_date:
                parse_date(to_date)
        except ValueError:
            raise ValueError("Dates must be in YYYY-MM-DD format.")

        if from_date and to_date and parse_date(from_date) > parse_date(to_date):
            raise ValueError("From Date cannot be later than To Date.")

        if report_kind == "boarders":
            rows = self.db.report_boarders()
            columns = [
                "Admission No.",
                "Learner Name",
                "Class",
                "Parent Name",
                "Parent Phone",
            ]
            data_rows = [
                [
                    row["admission_no"],
                    row["learner_name"],
                    row["class_level"],
                    row["parent_name"],
                    row["parent_phone"],
                ]
                for row in rows
            ]
            number_col, numbered_rows = add_number_column(data_rows)
            return "Boarders List", number_col + columns, numbered_rows, "", ""

        if report_kind == "school_bus":
            rows = self.db.report_school_bus_users()
            columns = ["Admission No.", "Learner Name", "Parent Name", "Parent Phone"]
            data_rows = [
                [row["admission_no"], row["learner_name"], row["parent_name"], row["parent_phone"]]
                for row in rows
            ]
            number_col, numbered_rows = add_number_column(data_rows)
            return "School Bus Users", number_col + columns, numbered_rows, "", ""

        if report_kind == "all_learners":
            rows = self.db.report_all_learners()
            columns = [
                "Admission No.",
                "Learner Name",
                "Class",
                "Parent Name",
                "Parent Phone",
                "Boarding Status",
                "Transport Mode",
            ]
            data_rows = [
                [
                    row["admission_no"],
                    row["learner_name"],
                    row["class_level"],
                    row["parent_name"],
                    row["parent_phone"],
                    row["boarding_status"],
                    row["transport_mode"],
                ]
                for row in rows
            ]
            number_col, numbered_rows = add_number_column(data_rows)
            return "All Learners with Parent Details", number_col + columns, numbered_rows, "", ""

        rows = self.db.report_parents_with_multiple_children()
        columns = ["Parent Name", "Parent Phone", "Count of Children", "Children"]
        data_rows = [
            [row["parent_name"], row["parent_phone"], str(row["child_count"]), row["children"]]
            for row in rows
        ]
        if report_kind == "multi_parents":
            number_col, numbered_rows = add_number_column(data_rows)
            return "Parents with Multiple Learners", number_col + columns, numbered_rows, "", ""

        rows = self.db.report_reported_learners(from_date, to_date)
        columns = [
            "Report Date",
            "Report Time",
            "Admission No.",
            "Learner Name",
            "Class",
            "Boarding Status",
            "Arrival Transport",
            "Accompanied By",
            "Accompanied Phone",
        ]
        data_rows = [
            [
                row["report_date"],
                row["report_time"],
                row["admission_no"],
                row["learner_name"],
                row["class_level"],
                row["boarding_status"],
                row["arrival_transport_mode"],
                row["accompanied_by_name"],
                row["accompanied_by_phone"],
            ]
            for row in rows
        ]
        number_col, numbered_rows = add_number_column(data_rows)
        return "Reported Learners", number_col + columns, numbered_rows, from_date, to_date

    def _build_preview_text(
        self,
        title: str,
        columns: list[str],
        rows: list[list[str]],
        from_date: str,
        to_date: str,
    ) -> str:
        max_width = 28
        widths = []
        for idx, col in enumerate(columns):
            longest = len(col)
            for row in rows:
                longest = max(longest, len(str(row[idx])))
            widths.append(min(longest, max_width))

        def fit(value: str, width: int) -> str:
            if len(value) <= width:
                return value.ljust(width)
            return f"{value[: width - 3]}..."

        lines = [self.setting_school_name_var.get().strip() or "Tumaini Academy", title]
        lines.append(f"Generated: {datetime.now():%Y-%m-%d %H:%M:%S}")
        if from_date and to_date:
            lines.append(f"Date range: {from_date} to {to_date}")
        lines.append(f"Total records: {len(rows)}")
        lines.append("")
        lines.append(" | ".join(fit(col, widths[i]) for i, col in enumerate(columns)))
        lines.append("-" * min(160, sum(widths) + (3 * (len(columns) - 1))))
        if rows:
            for row in rows:
                lines.append(" | ".join(fit(str(row[i]), widths[i]) for i in range(len(columns))))
        else:
            lines.append("No data for this selection.")
        return "\n".join(lines)

    def generate_report(self) -> None:
        try:
            title, columns, rows, from_date, to_date = self._prepare_report_data()
        except ValueError as error:
            messagebox.showerror("Report Error", str(error))
            return

        output_dir = Path(self.setting_report_dir_var.get().strip() or str(DEFAULT_REPORT_DIR))
        output_dir.mkdir(parents=True, exist_ok=True)
        fmt = self.report_format_var.get()
        slug_map = {
            "Boarders List": "BoardersList",
            "School Bus Users": "SchoolBusUsers",
            "All Learners with Parent Details": "AllLearnersWithParents",
            "Parents with Multiple Learners": "ParentsWithMultipleLearners",
            "Reported Learners": "ReportedLearners",
        }
        filename = f"{slug_map[title]}_{report_timestamp()}"
        extension = "pdf" if fmt == "PDF" else "docx"
        output_file = output_dir / f"{filename}.{extension}"
        orientation_preference = self.report_orientation_var.get().strip().title()
        if orientation_preference not in REPORT_ORIENTATIONS:
            orientation_preference = "Auto"

        exporter = ReportExporter(self.setting_school_name_var.get().strip() or "Tumaini Academy")
        try:
            if fmt == "PDF":
                used_orientation = exporter.export_pdf(
                    title,
                    columns,
                    rows,
                    output_file,
                    from_date,
                    to_date,
                    orientation_preference,
                    self.setting_pdf_template_portrait_var.get().strip(),
                    self.setting_pdf_template_landscape_var.get().strip(),
                )
            else:
                used_orientation = exporter.export_docx(
                    title,
                    columns,
                    rows,
                    output_file,
                    from_date,
                    to_date,
                    orientation_preference,
                    self.setting_word_template_portrait_var.get().strip(),
                    self.setting_word_template_landscape_var.get().strip(),
                )
        except RuntimeError as err:
            messagebox.showerror("Report Export Error", str(err))
            return

        self.last_generated_report = output_file
        preview = self._build_preview_text(title, columns, rows, from_date, to_date)
        self.report_preview.configure(state=tk.NORMAL)
        self.report_preview.delete("1.0", tk.END)
        self.report_preview.insert("1.0", preview)
        self.report_preview.configure(state=tk.DISABLED)

        self._set_status(f"Report generated ({used_orientation}): {output_file}")
        messagebox.showinfo(
            "Report Generated",
            f"Report saved to:\n{output_file}\n\nOrientation used: {used_orientation}",
        )
    def open_last_report(self) -> None:
        if not self.last_generated_report:
            messagebox.showinfo("Open Report", "No report has been generated in this session.")
            return
        if not self.last_generated_report.exists():
            messagebox.showerror("Open Report", "Last generated report file cannot be found.")
            return
        open_path(self.last_generated_report)

    def open_reports_folder(self) -> None:
        folder = Path(self.setting_report_dir_var.get().strip() or str(DEFAULT_REPORT_DIR))
        folder.mkdir(parents=True, exist_ok=True)
        open_path(folder)

    def backup_database(self) -> None:
        destination = filedialog.asksaveasfilename(
            title="Backup Database",
            defaultextension=".db",
            initialfile=f"tumaini_backup_{report_timestamp()}.db",
            filetypes=[("SQLite Database", "*.db"), ("All Files", "*.*")],
        )
        if not destination:
            return
        self.db.backup_database(Path(destination))
        self._set_status(f"Database backup saved: {destination}")
        messagebox.showinfo("Backup Complete", f"Database backup saved to:\n{destination}")

    def restore_database(self) -> None:
        source = filedialog.askopenfilename(
            title="Select Backup Database",
            filetypes=[("SQLite Database", "*.db"), ("All Files", "*.*")],
        )
        if not source:
            return
        confirm = messagebox.askyesno(
            "Confirm Restore",
            "Restore will overwrite the current database.\nDo you want to continue?",
        )
        if not confirm:
            return
        try:
            self.db.restore_database(Path(source))
            self.db.initialize()
        except Exception as error:
            messagebox.showerror("Restore Failed", str(error))
            return
        self._load_initial_data()
        self._set_status(f"Database restored from {source}")
        messagebox.showinfo("Restore Complete", "Database restored successfully.")

    def browse_report_folder(self) -> None:
        selected = filedialog.askdirectory(title="Select Default Report Folder")
        if selected:
            self.setting_report_dir_var.set(selected)

    def browse_logo_file(self) -> None:
        selected = filedialog.askopenfilename(
            title="Select School Logo Image",
            filetypes=[("PNG Image", "*.png"), ("All Files", "*.*")],
        )
        if selected:
            self.setting_logo_var.set(selected)

    def launch_exam_app(self) -> None:
        exam_file = BASE_DIR / "exam_app.py"
        if not exam_file.exists():
            messagebox.showerror("Exam App", f"Exam app file not found:\n{exam_file}")
            return
        subprocess.Popen([sys.executable, str(exam_file)])
        self._set_status("Exam marks desktop app launched.")

    def launch_exam_portal(self) -> None:
        portal_file = BASE_DIR / "exam_portal.py"
        if not portal_file.exists():
            messagebox.showerror("Exam Portal", f"Exam portal file not found:\n{portal_file}")
            return
        subprocess.Popen([sys.executable, str(portal_file)])
        self._set_status("Exam portal started.")

    def open_exam_portal_url(self) -> None:
        url = self.exam_portal_url_var.get().strip() or "http://127.0.0.1:5050"
        webbrowser.open(url)
        self._set_status(f"Opened portal URL: {url}")

    def save_exam_portal_url(self) -> None:
        url = self.exam_portal_url_var.get().strip()
        if not url:
            messagebox.showerror("Exam URL", "Portal URL cannot be empty.")
            return
        self.db.set_setting("exam_base_url", url)
        self._set_status("Exam portal URL saved.")
        messagebox.showinfo("Exam URL", "Portal URL saved.")

    def browse_template_file(
        self,
        target_var: tk.StringVar,
        title: str,
        filetypes: list[tuple[str, str]],
    ) -> None:
        selected = filedialog.askopenfilename(title=title, filetypes=filetypes)
        if selected:
            target_var.set(selected)

    def save_settings(self) -> None:
        school_name = self.setting_school_name_var.get().strip()
        report_dir = self.setting_report_dir_var.get().strip()
        password = self.setting_password_var.get()
        logo_path = self.setting_logo_var.get().strip()
        orientation = self.report_orientation_var.get().strip().title()
        if orientation not in REPORT_ORIENTATIONS:
            orientation = "Auto"
        word_template_portrait = self.setting_word_template_portrait_var.get().strip()
        word_template_landscape = self.setting_word_template_landscape_var.get().strip()
        pdf_template_portrait = self.setting_pdf_template_portrait_var.get().strip()
        pdf_template_landscape = self.setting_pdf_template_landscape_var.get().strip()
        exam_base_url = self.exam_portal_url_var.get().strip()

        if not school_name:
            messagebox.showerror("Settings", "School name cannot be empty.")
            return
        if not report_dir:
            messagebox.showerror("Settings", "Default report folder cannot be empty.")
            return

        report_path = Path(report_dir)
        report_path.mkdir(parents=True, exist_ok=True)
        self.db.set_setting("school_name", school_name)
        self.db.set_setting("default_report_dir", str(report_path))
        self.db.set_setting("app_password", password)
        self.db.set_setting("logo_path", logo_path)
        self.db.set_setting("report_orientation", orientation)
        self.db.set_setting("word_template_portrait", word_template_portrait)
        self.db.set_setting("word_template_landscape", word_template_landscape)
        self.db.set_setting("pdf_template_portrait", pdf_template_portrait)
        self.db.set_setting("pdf_template_landscape", pdf_template_landscape)
        self.db.set_setting("exam_base_url", exam_base_url or "http://127.0.0.1:5050")

        self.school_name = school_name
        self.default_report_dir = report_path
        self.logo_path = Path(logo_path) if logo_path else DEFAULT_LOGO_PATH
        self.report_orientation_var.set(orientation)
        self._load_logo()
        self._set_status("Settings saved.")
        messagebox.showinfo("Settings", "Settings saved successfully.")

    def open_user_manual(self) -> None:
        if MANUAL_PDF.exists():
            open_path(MANUAL_PDF)
            return
        if MANUAL_MD.exists():
            open_path(MANUAL_MD)
            return
        messagebox.showinfo("User Manual", "User manual file was not found in the docs folder.")

    def show_about(self) -> None:
        text = (
            f"{APP_TITLE}\n\n"
            "Offline desktop application for learner management,\n"
            "daily reporting, report generation, and backup/restore."
        )
        messagebox.showinfo("About", text)

    def _exit_app(self) -> None:
        self.destroy()


def main() -> None:
    app = TumainiApp()
    app.mainloop()


if __name__ == "__main__":
    main()
