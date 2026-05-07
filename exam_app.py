from __future__ import annotations

import subprocess
import sys
import webbrowser
from datetime import datetime
from pathlib import Path
import tkinter as tk
from tkinter import messagebox, ttk

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT

    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

from exam_shared import (
    CLASS_LEVELS,
    DEFAULT_REPORT_DIR,
    DEFAULT_TERMS,
    ExamDatabase,
    compute_level,
    report_timestamp,
)

APP_TITLE = "Tumaini Academy Exam Marks and Report System"
ORIENTATION_OPTIONS = ("Auto", "Portrait", "Landscape")


class MarksExporter:
    def __init__(self, school_name: str) -> None:
        self.school_name = school_name

    def _resolve_orientation(self, orientation: str, columns: list[str], rows: list[list[str]]) -> str:
        value = (orientation or "Auto").strip().title()
        if value in ("Portrait", "Landscape"):
            return value
        if len(columns) >= 8:
            return "Landscape"
        for row in rows[:50]:
            if any(len(str(item)) > 28 for item in row):
                return "Landscape"
        return "Portrait"

    def export_pdf(
        self,
        title: str,
        columns: list[str],
        rows: list[list[str]],
        output_file: Path,
        orientation: str = "Auto",
    ) -> str:
        if not REPORTLAB_AVAILABLE:
            raise RuntimeError("PDF export requires reportlab (install requirements.txt).")

        resolved = self._resolve_orientation(orientation, columns, rows)
        doc = SimpleDocTemplate(str(output_file), pagesize=landscape(A4) if resolved == "Landscape" else A4)
        styles = getSampleStyleSheet()
        content: list[object] = [
            Paragraph(self.school_name, styles["Title"]),
            Paragraph(title, styles["Heading2"]),
            Paragraph(f"Generated: {datetime.now():%Y-%m-%d %H:%M:%S}", styles["Normal"]),
            Paragraph(f"Total records: {len(rows)}", styles["Normal"]),
            Spacer(1, 10),
        ]

        table_rows = [columns] + rows if rows else [columns, ["No data for this selection."] + [""] * (len(columns) - 1)]
        table = Table(table_rows, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F766E")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
                ]
            )
        )
        content.append(table)
        doc.build(content)
        return resolved

    def export_docx(
        self,
        title: str,
        columns: list[str],
        rows: list[list[str]],
        output_file: Path,
        orientation: str = "Auto",
    ) -> str:
        if not DOCX_AVAILABLE:
            raise RuntimeError("Word export requires python-docx (install requirements.txt).")

        resolved = self._resolve_orientation(orientation, columns, rows)
        doc = Document()
        for section in doc.sections:
            if resolved == "Landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                if section.page_width < section.page_height:
                    section.page_width, section.page_height = section.page_height, section.page_width
            else:
                section.orientation = WD_ORIENT.PORTRAIT
                if section.page_width > section.page_height:
                    section.page_width, section.page_height = section.page_height, section.page_width

        doc.add_heading(self.school_name, 0)
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"Generated: {datetime.now():%Y-%m-%d %H:%M:%S}")
        doc.add_paragraph(f"Total records: {len(rows)}")
        if rows:
            table = doc.add_table(rows=1, cols=len(columns))
            table.style = "Table Grid"
            for idx, header in enumerate(columns):
                table.rows[0].cells[idx].text = header
            for row in rows:
                cells = table.add_row().cells
                for idx, value in enumerate(row):
                    cells[idx].text = str(value)
        else:
            doc.add_paragraph("No data for this selection.")
        doc.save(output_file)
        return resolved


class ExamApp(tk.Tk):
    def __init__(self) -> None:
        super().__init__()
        self.db = ExamDatabase()
        self.db.initialize()
        self.base_dir = Path(__file__).resolve().parent
        self.selected_form_id: int | None = None
        self.selected_admission_no: str | None = None
        self.form_label_to_id: dict[str, int] = {}

        self.title(APP_TITLE)
        self.geometry("1220x780")
        self.minsize(1080, 700)

        self._build_vars()
        self._build_ui()
        self.refresh_learning_areas()
        self.refresh_forms()
        self.load_marks()

    def _build_vars(self) -> None:
        current_year = datetime.now().year
        self.status_var = tk.StringVar(value="Ready")

        self.exam_name_var = tk.StringVar(value="Mid Term")
        self.term_var = tk.StringVar(value=DEFAULT_TERMS[0])
        self.year_var = tk.StringVar(value=str(current_year))
        self.class_var = tk.StringVar(value=CLASS_LEVELS[0])
        self.subject_var = tk.StringVar(value="Mathematics")
        self.max_marks_var = tk.StringVar(value="100")
        self.learning_area_var = tk.StringVar()
        self.learning_area_id_map: dict[str, int] = {}
        self.link_var = tk.StringVar()

        self.form_choice_var = tk.StringVar()
        self.report_format_var = tk.StringVar(value="PDF")
        self.report_orientation_var = tk.StringVar(value="Auto")
        self.teacher_name_var = tk.StringVar(value="Office Entry")
        self.manual_mark_var = tk.StringVar()

        self.base_url_var = tk.StringVar(value=self.db.get_setting("exam_base_url", "http://127.0.0.1:5050"))
        self.area_class_filter_var = tk.StringVar(value=CLASS_LEVELS[0])
        self.area_name_var = tk.StringVar()
        self.area_max_var = tk.StringVar(value="100")
        self.area_min_var = tk.StringVar(value="0")
        self.area_formula_var = tk.StringVar(value="80:Exceeds,60:Meets,40:Approaching,0:Below")
        self.selected_area_id: int | None = None

    def _build_ui(self) -> None:
        root = ttk.Frame(self, padding=8)
        root.pack(fill=tk.BOTH, expand=True)

        ttk.Label(root, text=APP_TITLE, font=("Segoe UI", 12, "bold")).pack(anchor="w", pady=(0, 8))

        notebook = ttk.Notebook(root)
        notebook.pack(fill=tk.BOTH, expand=True)

        self.forms_tab = ttk.Frame(notebook, padding=10)
        self.marks_tab = ttk.Frame(notebook, padding=10)
        self.settings_tab = ttk.Frame(notebook, padding=10)
        notebook.add(self.forms_tab, text="Online Forms")
        notebook.add(self.marks_tab, text="Marks & Reports")
        notebook.add(self.settings_tab, text="Settings")

        self._build_forms_tab()
        self._build_marks_tab()
        self._build_learning_areas_tab()
        self._build_settings_tab()

        ttk.Label(self, textvariable=self.status_var, anchor="w").pack(fill=tk.X, padx=8, pady=(0, 8))

    def _build_forms_tab(self) -> None:
        create_box = ttk.LabelFrame(self.forms_tab, text="Create Teacher Link", padding=10)
        create_box.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(create_box, text="Exam Name").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(create_box, textvariable=self.exam_name_var, width=28).grid(row=0, column=1, sticky="w", padx=5, pady=5)

        ttk.Label(create_box, text="Term").grid(row=0, column=2, sticky="w", padx=5, pady=5)
        ttk.Combobox(create_box, textvariable=self.term_var, values=DEFAULT_TERMS, state="readonly", width=14).grid(
            row=0, column=3, sticky="w", padx=5, pady=5
        )

        ttk.Label(create_box, text="Year").grid(row=0, column=4, sticky="w", padx=5, pady=5)
        ttk.Entry(create_box, textvariable=self.year_var, width=10).grid(row=0, column=5, sticky="w", padx=5, pady=5)

        ttk.Label(create_box, text="Class").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        class_combo = ttk.Combobox(create_box, textvariable=self.class_var, values=CLASS_LEVELS, state="readonly", width=24)
        class_combo.grid(row=1, column=1, sticky="w", padx=5, pady=5)
        class_combo.bind("<<ComboboxSelected>>", lambda _e: self.refresh_learning_areas())

        ttk.Label(create_box, text="Subject").grid(row=1, column=2, sticky="w", padx=5, pady=5)
        ttk.Entry(create_box, textvariable=self.subject_var, width=20).grid(row=1, column=3, sticky="w", padx=5, pady=5)

        ttk.Label(create_box, text="Max Marks").grid(row=1, column=4, sticky="w", padx=5, pady=5)
        ttk.Entry(create_box, textvariable=self.max_marks_var, width=10).grid(row=1, column=5, sticky="w", padx=5, pady=5)

        ttk.Label(create_box, text="Learning Area (optional)").grid(row=2, column=0, sticky="w", padx=5, pady=5)
        self.learning_area_combo = ttk.Combobox(
            create_box,
            textvariable=self.learning_area_var,
            state="readonly",
            width=46,
        )
        self.learning_area_combo.grid(row=2, column=1, columnspan=3, sticky="we", padx=5, pady=5)
        self.learning_area_combo.bind("<<ComboboxSelected>>", self.on_learning_area_selected)

        ttk.Button(create_box, text="Generate Link", command=self.create_link).grid(
            row=2, column=0, sticky="w", padx=5, pady=(8, 4)
        )

        ttk.Label(create_box, text="Generated Link").grid(row=3, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(create_box, textvariable=self.link_var, width=95).grid(
            row=3, column=1, columnspan=4, sticky="we", padx=5, pady=5
        )
        ttk.Button(create_box, text="Copy Link", command=self.copy_generated_link).grid(row=3, column=5, sticky="w", padx=5, pady=5)
        create_box.columnconfigure(1, weight=1)

        table_box = ttk.LabelFrame(self.forms_tab, text="Active Mark Entry Links", padding=10)
        table_box.pack(fill=tk.BOTH, expand=True)

        cols = ("no", "id", "created_at", "exam_name", "term", "year", "class_level", "subject", "max_marks", "token")
        self.forms_tree = ttk.Treeview(table_box, columns=cols, show="headings", height=14)
        headers = {
            "no": "No.",
            "id": "ID",
            "created_at": "Created",
            "exam_name": "Exam",
            "term": "Term",
            "year": "Year",
            "class_level": "Class",
            "subject": "Subject",
            "max_marks": "Max",
            "token": "Token",
        }
        for col in cols:
            self.forms_tree.heading(col, text=headers[col])
            width = 70 if col in ("no", "id", "term", "year", "max_marks") else 160
            self.forms_tree.column(col, width=width, anchor="w")
        self.forms_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.forms_tree.bind("<<TreeviewSelect>>", self.on_form_select)

        sb = ttk.Scrollbar(table_box, orient=tk.VERTICAL, command=self.forms_tree.yview)
        self.forms_tree.configure(yscrollcommand=sb.set)
        sb.pack(side=tk.LEFT, fill=tk.Y)

        actions = ttk.Frame(self.forms_tab)
        actions.pack(fill=tk.X, pady=(8, 0))
        ttk.Button(actions, text="Refresh", command=self.refresh_forms).pack(side=tk.LEFT)
        ttk.Button(actions, text="Open Selected Link", command=self.open_selected_link).pack(side=tk.LEFT, padx=(8, 0))

    def _build_marks_tab(self) -> None:
        top = ttk.LabelFrame(self.marks_tab, text="Load Marks", padding=10)
        top.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(top, text="Exam Form").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        self.form_choice_combo = ttk.Combobox(top, textvariable=self.form_choice_var, state="readonly", width=70)
        self.form_choice_combo.grid(row=0, column=1, columnspan=4, sticky="we", padx=5, pady=5)
        ttk.Button(top, text="Load", command=self.load_marks).grid(row=0, column=5, sticky="w", padx=5, pady=5)

        ttk.Label(top, text="Teacher Name").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(top, textvariable=self.teacher_name_var, width=26).grid(row=1, column=1, sticky="w", padx=5, pady=5)

        ttk.Label(top, text="Selected Mark").grid(row=1, column=2, sticky="w", padx=5, pady=5)
        ttk.Entry(top, textvariable=self.manual_mark_var, width=12).grid(row=1, column=3, sticky="w", padx=5, pady=5)
        ttk.Button(top, text="Update Selected", command=self.update_selected_mark).grid(row=1, column=4, sticky="w", padx=5, pady=5)

        ttk.Label(top, text="Format").grid(row=2, column=0, sticky="w", padx=5, pady=5)
        ttk.Combobox(top, textvariable=self.report_format_var, values=("PDF", "Word"), state="readonly", width=12).grid(
            row=2, column=1, sticky="w", padx=5, pady=5
        )
        ttk.Label(top, text="Orientation").grid(row=2, column=2, sticky="w", padx=5, pady=5)
        ttk.Combobox(
            top,
            textvariable=self.report_orientation_var,
            values=ORIENTATION_OPTIONS,
            state="readonly",
            width=12,
        ).grid(row=2, column=3, sticky="w", padx=5, pady=5)
        ttk.Button(top, text="Download Report", command=self.download_report).grid(row=2, column=4, sticky="w", padx=5, pady=5)
        top.columnconfigure(1, weight=1)

        box = ttk.LabelFrame(self.marks_tab, text="Marks List", padding=10)
        box.pack(fill=tk.BOTH, expand=True)
        cols = ("no", "admission_no", "learner_name", "marks", "percent", "level", "teacher", "submitted_at")
        self.marks_tree = ttk.Treeview(box, columns=cols, show="headings", height=16)
        headers = {
            "no": "No.",
            "admission_no": "Admission No.",
            "learner_name": "Learner Name",
            "marks": "Marks",
            "percent": "%",
            "level": "Level",
            "teacher": "Submitted By",
            "submitted_at": "Submitted At",
        }
        for col in cols:
            self.marks_tree.heading(col, text=headers[col])
            width = 80 if col in ("no", "percent", "marks", "level") else 170
            self.marks_tree.column(col, width=width, anchor="w")
        self.marks_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.marks_tree.bind("<<TreeviewSelect>>", self.on_mark_select)
        sb = ttk.Scrollbar(box, orient=tk.VERTICAL, command=self.marks_tree.yview)
        self.marks_tree.configure(yscrollcommand=sb.set)
        sb.pack(side=tk.LEFT, fill=tk.Y)

    def _build_learning_areas_tab(self) -> None:
        tab = self.exams_tab_content
        controls = ttk.LabelFrame(tab, text="Learning Areas per Class", padding=10)
        controls.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(controls, text="Class").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        ttk.Combobox(
            controls,
            textvariable=self.area_class_filter_var,
            values=CLASS_LEVELS,
            state="readonly",
            width=22,
        ).grid(row=0, column=1, sticky="w", padx=5, pady=5)
        ttk.Button(controls, text="Load", command=self.refresh_learning_areas).grid(row=0, column=2, sticky="w", padx=5, pady=5)

        ttk.Label(controls, text="Name").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(controls, textvariable=self.area_name_var, width=28).grid(row=1, column=1, sticky="w", padx=5, pady=5)

        ttk.Label(controls, text="Max Marks").grid(row=1, column=2, sticky="w", padx=5, pady=5)
        ttk.Entry(controls, textvariable=self.area_max_var, width=10).grid(row=1, column=3, sticky="w", padx=5, pady=5)

        ttk.Label(controls, text="Min Marks").grid(row=1, column=4, sticky="w", padx=5, pady=5)
        ttk.Entry(controls, textvariable=self.area_min_var, width=10).grid(row=1, column=5, sticky="w", padx=5, pady=5)

        ttk.Label(controls, text="Level Formula (e.g. 80:Exceeds,60:Meets,40:Approaching,0:Below)").grid(
            row=2, column=0, columnspan=6, sticky="w", padx=5, pady=5
        )
        ttk.Entry(controls, textvariable=self.area_formula_var, width=90).grid(
            row=3, column=0, columnspan=6, sticky="we", padx=5, pady=5
        )

        ttk.Button(controls, text="Save / Update", command=self.save_learning_area).grid(row=4, column=0, sticky="w", padx=5, pady=(8, 4))
        ttk.Button(controls, text="Clear", command=self.clear_learning_area_form).grid(row=4, column=1, sticky="w", padx=5, pady=(8, 4))
        ttk.Button(controls, text="Delete", command=self.delete_learning_area).grid(row=4, column=2, sticky="w", padx=5, pady=(8, 4))
        controls.columnconfigure(1, weight=1)

        table_box = ttk.LabelFrame(tab, text="Learning Areas List", padding=10)
        table_box.pack(fill=tk.BOTH, expand=True)
        cols = ("no", "id", "class_level", "name", "max_marks", "min_marks", "level_formula")
        self.area_tree = ttk.Treeview(table_box, columns=cols, show="headings", height=12)
        headers = {
            "no": "No.",
            "id": "ID",
            "class_level": "Class",
            "name": "Name",
            "max_marks": "Max",
            "min_marks": "Min",
            "level_formula": "Level Formula",
        }
        for col in cols:
            self.area_tree.heading(col, text=headers[col])
            width = 80 if col in ("no", "id", "max_marks", "min_marks") else 180
            self.area_tree.column(col, width=width, anchor="w")
        self.area_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.area_tree.bind("<<TreeviewSelect>>", self.on_area_select)
        sb = ttk.Scrollbar(table_box, orient=tk.VERTICAL, command=self.area_tree.yview)
        self.area_tree.configure(yscrollcommand=sb.set)
        sb.pack(side=tk.LEFT, fill=tk.Y)

    def _build_settings_tab(self) -> None:
        box = ttk.LabelFrame(self.settings_tab, text="Online Portal", padding=10)
        box.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(box, text="Base URL for Links").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(box, textvariable=self.base_url_var, width=60).grid(row=0, column=1, sticky="we", padx=5, pady=5)
        ttk.Button(box, text="Save", command=self.save_settings).grid(row=0, column=2, sticky="w", padx=5, pady=5)

        ttk.Button(box, text="Start Local Portal", command=self.start_local_portal).grid(
            row=1, column=0, sticky="w", padx=5, pady=(8, 5)
        )
        ttk.Label(
            box,
            text="Teachers can use links like: <BaseURL>/fill/<token>.",
        ).grid(row=1, column=1, columnspan=2, sticky="w", padx=5, pady=(8, 5))
        box.columnconfigure(1, weight=1)

    def _set_status(self, text: str) -> None:
        self.status_var.set(text)

    def refresh_forms(self) -> None:
        forms = self.db.list_exam_forms(active_only=True)
        self.forms_tree.delete(*self.forms_tree.get_children())
        self.form_label_to_id.clear()
        options: list[str] = []

        for idx, row in enumerate(forms, start=1):
            self.forms_tree.insert(
                "",
                tk.END,
                values=(
                    idx,
                    row["id"],
                    row["created_at"],
                    row["exam_name"],
                    row["term"],
                    row["year"],
                    row["class_level"],
                    row["subject"],
                    row["max_marks"],
                    row["token"],
                ),
            )
            label = (
                f"{row['id']} | {row['exam_name']} | {row['class_level']} | "
                f"{row['subject']} | {row['term']} {row['year']}"
            )
            options.append(label)
            self.form_label_to_id[label] = int(row["id"])

        self.form_choice_combo.configure(values=options)
        if options and self.form_choice_var.get() not in self.form_label_to_id:
            self.form_choice_var.set(options[0])
        self._set_status(f"Loaded {len(forms)} active link(s).")

    def create_link(self) -> None:
        exam_name = self.exam_name_var.get().strip()
        term = self.term_var.get().strip()
        year_text = self.year_var.get().strip()
        class_level = self.class_var.get().strip()
        subject = self.subject_var.get().strip()
        max_marks_text = self.max_marks_var.get().strip()

        if not exam_name or not subject:
            messagebox.showerror("Validation", "Exam name and subject are required.")
            return
        try:
            year = int(year_text)
        except ValueError:
            messagebox.showerror("Validation", "Year must be a number.")
            return
        try:
            max_marks = float(max_marks_text)
        except ValueError:
            messagebox.showerror("Validation", "Max marks must be numeric.")
            return

        learners = self.db.get_learners_by_class(class_level)
        if not learners:
            messagebox.showwarning("No Learners", f"No learners found in class '{class_level}'.")
            return

        area_id = None
        chosen_area = self.learning_area_var.get().strip()
        if chosen_area and chosen_area in self.learning_area_id_map:
            area_id = self.learning_area_id_map[chosen_area]

        try:
            token = self.db.create_exam_form(exam_name, term, year, class_level, subject, max_marks, area_id)
        except ValueError as err:
            messagebox.showerror("Create Link", str(err))
            return

        base = self.base_url_var.get().strip().rstrip("/") or "http://127.0.0.1:5050"
        link = f"{base}/fill/{token}"
        self.link_var.set(link)
        self.refresh_forms()
        self._set_status(f"Link generated for {exam_name} - {class_level} {subject}.")

    def copy_generated_link(self) -> None:
        link = self.link_var.get().strip()
        if not link:
            messagebox.showinfo("Copy Link", "No link generated yet.")
            return
        self.clipboard_clear()
        self.clipboard_append(link)
        self._set_status("Generated link copied.")
        messagebox.showinfo("Copy Link", "Link copied to clipboard.")

    def on_form_select(self, _event: tk.Event | None = None) -> None:
        selected = self.forms_tree.selection()
        if not selected:
            return
        values = self.forms_tree.item(selected[0], "values")
        if not values:
            return
        self.selected_form_id = int(values[1])
        token = str(values[9])
        base = self.base_url_var.get().strip().rstrip("/") or "http://127.0.0.1:5050"
        self.link_var.set(f"{base}/fill/{token}")
        for label, exam_id in self.form_label_to_id.items():
            if exam_id == self.selected_form_id:
                self.form_choice_var.set(label)
                break

    def open_selected_link(self) -> None:
        link = self.link_var.get().strip()
        if not link:
            messagebox.showinfo("Open Link", "Select a form or generate a link first.")
            return
        webbrowser.open(link)

    def on_learning_area_selected(self, _event: tk.Event | None = None) -> None:
        name = self.learning_area_var.get().strip()
        if not name or name not in self.learning_area_id_map:
            return
        area_id = self.learning_area_id_map[name]
        area = self.db.get_learning_area(area_id)
        if not area:
            return
        # Prefill subject and max marks from the learning area but allow edits.
        self.subject_var.set(area["name"])
        self.max_marks_var.set(str(area["max_marks"]))

    def _selected_form_id_from_choice(self) -> int | None:
        choice = self.form_choice_var.get().strip()
        if choice and choice in self.form_label_to_id:
            return self.form_label_to_id[choice]
        return self.selected_form_id

    def load_marks(self) -> None:
        exam_id = self._selected_form_id_from_choice()
        self.marks_tree.delete(*self.marks_tree.get_children())
        self.selected_admission_no = None
        self.manual_mark_var.set("")

        if exam_id is None:
            self._set_status("Select an exam form to load marks.")
            return

        rows = self.db.get_marks_for_exam(exam_id)
        for idx, row in enumerate(rows, start=1):
            marks = "" if row["marks"] is None else float(row["marks"])
            percent = ""
            level = ""
            if row["marks"] is not None and float(row["max_marks"]) > 0:
                percent_value = (float(row["marks"]) / float(row["max_marks"])) * 100
                percent = f"{percent_value:.1f}"
                level = compute_level(float(row["marks"]), float(row["max_marks"]), row["level_formula"] or "")
            self.marks_tree.insert(
                "",
                tk.END,
                values=(
                    idx,
                    row["admission_no"],
                    row["learner_name"],
                    marks,
                    percent,
                    level,
                    row["teacher_name"] or "",
                    row["submitted_at"] or "",
                ),
            )
        self._set_status(f"Loaded marks for {len(rows)} learner(s).")

    def on_mark_select(self, _event: tk.Event | None = None) -> None:
        selected = self.marks_tree.selection()
        if not selected:
            return
        values = self.marks_tree.item(selected[0], "values")
        if not values:
            return
        self.selected_admission_no = str(values[1])
        self.manual_mark_var.set("" if values[3] == "" else str(values[3]))

    def update_selected_mark(self) -> None:
        exam_id = self._selected_form_id_from_choice()
        if exam_id is None:
            messagebox.showwarning("Update Mark", "Load an exam form first.")
            return
        if not self.selected_admission_no:
            messagebox.showwarning("Update Mark", "Select a learner from the marks list.")
            return
        mark_text = self.manual_mark_var.get().strip()
        if mark_text == "":
            messagebox.showwarning("Update Mark", "Enter marks.")
            return
        try:
            mark_value = float(mark_text)
        except ValueError:
            messagebox.showerror("Update Mark", "Marks must be numeric.")
            return

        teacher = self.teacher_name_var.get().strip() or "Office Entry"
        try:
            self.db.save_marks_for_exam(exam_id, {self.selected_admission_no: mark_value}, teacher_name=teacher)
        except ValueError as err:
            messagebox.showerror("Update Mark", str(err))
            return
        self.load_marks()
        self._set_status(f"Marks updated for {self.selected_admission_no}.")

    def download_report(self) -> None:
        exam_id = self._selected_form_id_from_choice()
        if exam_id is None:
            messagebox.showwarning("Download Report", "Select an exam form first.")
            return

        form = self.db.get_exam_form(exam_id)
        if not form:
            messagebox.showerror("Download Report", "Selected exam form was not found.")
            return

        marks_rows = self.db.get_marks_for_exam(exam_id)
        columns = [
            "No.",
            "Admission No.",
            "Learner Name",
            "Marks",
            "Percent",
            "Level",
            "Submitted By",
            "Submitted At",
        ]
        data: list[list[str]] = []
        for idx, row in enumerate(marks_rows, start=1):
            mark = "" if row["marks"] is None else f"{float(row['marks']):.2f}"
            percent = ""
            level = ""
            if row["marks"] is not None and float(row["max_marks"]) > 0:
                percent = f"{(float(row['marks']) / float(row['max_marks'])) * 100:.1f}%"
                level = compute_level(float(row["marks"]), float(row["max_marks"]), row["level_formula"] or "")
            data.append(
                [
                    str(idx),
                    str(row["admission_no"]),
                    str(row["learner_name"]),
                    mark,
                    percent,
                    level,
                    str(row["teacher_name"] or ""),
                    str(row["submitted_at"] or ""),
                ]
            )

        school_name = self.db.get_setting("school_name", "Tumaini Academy")
        title = (
            f"{form['exam_name']} - {form['term']} {form['year']} | "
            f"{form['class_level']} | {form['subject']} | Max {form['max_marks']}"
        )
        output_dir = DEFAULT_REPORT_DIR
        output_dir.mkdir(parents=True, exist_ok=True)
        slug = f"ExamMarks_{form['exam_name']}_{form['class_level']}_{form['subject']}_{report_timestamp()}".replace(" ", "_")
        fmt = self.report_format_var.get()
        ext = "pdf" if fmt == "PDF" else "docx"
        output_file = output_dir / f"{slug}.{ext}"

        exporter = MarksExporter(school_name)
        orientation = self.report_orientation_var.get()
        try:
            if fmt == "PDF":
                used = exporter.export_pdf(title, columns, data, output_file, orientation=orientation)
            else:
                used = exporter.export_docx(title, columns, data, output_file, orientation=orientation)
        except RuntimeError as err:
            messagebox.showerror("Download Report", str(err))
            return

        self._set_status(f"Report downloaded ({used}): {output_file}")
        messagebox.showinfo("Download Report", f"Saved report to:\n{output_file}\n\nOrientation: {used}")

    def save_settings(self) -> None:
        base_url = self.base_url_var.get().strip()
        if not base_url:
            messagebox.showerror("Settings", "Base URL cannot be empty.")
            return
        self.db.set_setting("exam_base_url", base_url)
        self._set_status("Settings saved.")
        messagebox.showinfo("Settings", "Base URL saved.")

    def start_local_portal(self) -> None:
        portal_file = self.base_dir / "exam_portal.py"
        if not portal_file.exists():
            messagebox.showerror("Portal", f"Portal file not found:\n{portal_file}")
            return
        subprocess.Popen([sys.executable, str(portal_file)])
        self._set_status("Local portal started on http://127.0.0.1:5050")
        messagebox.showinfo("Portal", "Local portal started.\n\nUse http://127.0.0.1:5050")

    # Learning Areas helpers
    def refresh_learning_areas(self) -> None:
        class_level = self.area_class_filter_var.get().strip()
        rows = self.db.list_learning_areas(class_level)
        self.area_tree.delete(*self.area_tree.get_children())
        for idx, row in enumerate(rows, start=1):
            self.area_tree.insert(
                "",
                tk.END,
                values=(
                    idx,
                    row["id"],
                    row["class_level"],
                    row["name"],
                    row["max_marks"],
                    row["min_marks"],
                    row["level_formula"],
                ),
            )
        # refresh combo map for form creation
        areas = self.db.list_learning_areas(self.class_var.get().strip())
        self.learning_area_id_map = {f"{a['name']} (Max {a['max_marks']})": a["id"] for a in areas}
        self.learning_area_combo.configure(values=list(self.learning_area_id_map.keys()))
        self._set_status(f"Loaded {len(rows)} learning areas for {class_level}.")

    def on_area_select(self, _event: tk.Event | None = None) -> None:
        selected = self.area_tree.selection()
        if not selected:
            return
        vals = self.area_tree.item(selected[0], "values")
        if not vals:
            return
        self.selected_area_id = int(vals[1])
        self.area_class_filter_var.set(str(vals[2]))
        self.area_name_var.set(str(vals[3]))
        self.area_max_var.set(str(vals[4]))
        self.area_min_var.set(str(vals[5]))
        self.area_formula_var.set(str(vals[6]))

    def clear_learning_area_form(self) -> None:
        self.selected_area_id = None
        self.area_name_var.set("")
        self.area_max_var.set("100")
        self.area_min_var.set("0")
        self.area_formula_var.set("80:Exceeds,60:Meets,40:Approaching,0:Below")

    def save_learning_area(self) -> None:
        try:
            max_marks = float(self.area_max_var.get().strip() or "0")
            min_marks = float(self.area_min_var.get().strip() or "0")
        except ValueError:
            messagebox.showerror("Learning Area", "Max/Min must be numeric.")
            return
        name = self.area_name_var.get().strip()
        if not name:
            messagebox.showerror("Learning Area", "Name is required.")
            return
        try:
            area_id = self.db.upsert_learning_area(
                self.selected_area_id,
                self.area_class_filter_var.get().strip(),
                name,
                max_marks,
                min_marks,
                self.area_formula_var.get().strip(),
            )
        except ValueError as err:
            messagebox.showerror("Learning Area", str(err))
            return
        self.selected_area_id = area_id
        self.refresh_learning_areas()
        self._set_status("Learning area saved.")

    def delete_learning_area(self) -> None:
        if not self.selected_area_id:
            messagebox.showwarning("Learning Area", "Select a learning area to delete.")
            return
        if not messagebox.askyesno("Confirm Delete", "Delete selected learning area?"):
            return
        self.db.delete_learning_area(self.selected_area_id)
        self.selected_area_id = None
        self.clear_learning_area_form()
        self.refresh_learning_areas()
        self._set_status("Learning area deleted.")


def main() -> None:
    app = ExamApp()
    app.mainloop()


if __name__ == "__main__":
    main()
